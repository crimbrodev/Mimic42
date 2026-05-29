from __future__ import annotations

import base64
import json
from datetime import datetime, timedelta
from typing import Any, Literal, Protocol
from uuid import UUID

from langchain_core.tools import BaseTool, StructuredTool
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker
from telethon import functions, types
from telethon.extensions import markdown


class TelethonRequestClient(Protocol):
    async def __call__(self, request: object) -> object: ...
    async def connect(self) -> None: ...
    async def disconnect(self) -> None: ...
    async def is_user_authorized(self) -> bool: ...
    async def send_message(self, entity: Any, message: str, **kwargs: Any) -> Any: ...
    async def send_file(self, entity: Any, file: Any, **kwargs: Any) -> Any: ...
    async def edit_message(self, entity: Any, message: Any, text: str, **kwargs: Any) -> Any: ...
    async def delete_messages(self, entity: Any, message_ids: list[int], **kwargs: Any) -> Any: ...
    async def forward_messages(
        self, entity: Any, messages: list[int], from_peer: Any, **kwargs: Any
    ) -> Any: ...
    async def pin_message(self, entity: Any, message: Any, **kwargs: Any) -> Any: ...
    async def unpin_message(self, entity: Any, message: Any, **kwargs: Any) -> Any: ...
    async def get_entity(self, entity: Any) -> Any: ...
    async def get_input_entity(self, entity: Any) -> Any: ...
    async def get_messages(self, entity: Any, **kwargs: Any) -> Any: ...
    async def get_dialogs(self, **kwargs: Any) -> Any: ...
    async def get_permissions(self, entity: Any, user: Any) -> Any: ...
    async def edit_permissions(self, entity: Any, user: Any, **kwargs: Any) -> Any: ...
    async def edit_admin(self, entity: Any, user: Any, **kwargs: Any) -> Any: ...
    async def kick_participant(self, entity: Any, user: Any) -> Any: ...
    async def upload_file(self, file: Any, **kwargs: Any) -> Any: ...
    async def download_media(self, message: Any, file: Any = None, **kwargs: Any) -> Any: ...
    def iter_messages(self, entity: Any, **kwargs: Any) -> Any: ...
    def iter_participants(self, entity: Any, **kwargs: Any) -> Any: ...
    def iter_admin_log(self, entity: Any, **kwargs: Any) -> Any: ...


class CustomMarkdown:
    """Telethon parse mode to format premium custom emoji and markdown formatting."""

    @staticmethod
    def parse(text: str) -> tuple[str, list[types.TypeMessageEntity]]:
        if not text:
            return "", []
        parsed_text, entities = markdown.parse(text)
        new_entities = []
        for entity in entities:
            if isinstance(entity, types.MessageEntityTextUrl) and entity.url.startswith("emoji/"):
                try:
                    doc_id = int(entity.url.split("/")[1])
                    new_entities.append(
                        types.MessageEntityCustomEmoji(
                            offset=entity.offset,
                            length=entity.length,
                            document_id=doc_id,
                        )
                    )
                except (ValueError, IndexError):
                    new_entities.append(entity)
            else:
                new_entities.append(entity)
        return parsed_text, new_entities

    @staticmethod
    def unparse(text: str, entities: list[types.TypeMessageEntity]) -> str:
        if not text:
            return ""
        if not entities:
            return markdown.unparse(text, [])

        temp_entities = []
        for entity in entities:
            if isinstance(entity, types.MessageEntityCustomEmoji):
                temp_entities.append(
                    types.MessageEntityTextUrl(
                        offset=entity.offset,
                        length=entity.length,
                        url=f"emoji/{entity.document_id}",
                    )
                )
            else:
                temp_entities.append(entity)
        return markdown.unparse(text, temp_entities)


def parse_media_id(media_id: str) -> tuple[str, int, int, bytes, int]:
    """Parse media ID in format type:id:access_hash:file_reference_hex:dc_id."""
    parts = media_id.split(":")
    if len(parts) < 5:
        raise ValueError(f"Invalid media ID format: {media_id}")
    media_type = parts[0]
    obj_id = int(parts[1])
    access_hash = int(parts[2])
    file_reference = bytes.fromhex(parts[3])
    dc_id = int(parts[4])
    return media_type, obj_id, access_hash, file_reference, dc_id


def format_media_object(msg: Any) -> str | None:
    """Format message media to serialized media ID string."""
    if not msg or not getattr(msg, "media", None):
        return None
    media = msg.media

    if isinstance(media, types.MessageMediaPhoto) and media.photo:
        photo = media.photo
        if isinstance(photo, types.Photo):
            ref_hex = photo.file_reference.hex() if photo.file_reference else ""
            return f"photo:{photo.id}:{photo.access_hash}:{ref_hex}:{photo.dc_id}"

    elif isinstance(media, types.MessageMediaDocument) and media.document:
        doc = media.document
        if isinstance(doc, types.Document):
            ref_hex = doc.file_reference.hex() if doc.file_reference else ""

            # Check for sticker attribute
            sticker_attr = next(
                (
                    attr
                    for attr in doc.attributes
                    if isinstance(attr, types.DocumentAttributeSticker)
                ),
                None,
            )
            if sticker_attr:
                emoji = sticker_attr.alt or ""
                pack_name = ""
                if isinstance(sticker_attr.stickerset, types.InputStickerSetShortName):
                    pack_name = sticker_attr.stickerset.short_name
                elif isinstance(sticker_attr.stickerset, types.InputStickerSetID):
                    pack_name = str(sticker_attr.stickerset.id)
                return (
                    f"sticker:{doc.id}:{doc.access_hash}:{ref_hex}:{doc.dc_id}:{emoji}:{pack_name}"
                )

            # Check for voice note
            audio_attr = next(
                (attr for attr in doc.attributes if isinstance(attr, types.DocumentAttributeAudio)),
                None,
            )
            if audio_attr and audio_attr.voice:
                return f"voice:{doc.id}:{doc.access_hash}:{ref_hex}:{doc.dc_id}"

            # Check for round video note
            video_attr = next(
                (attr for attr in doc.attributes if isinstance(attr, types.DocumentAttributeVideo)),
                None,
            )
            if video_attr and video_attr.round_message:
                return f"round:{doc.id}:{doc.access_hash}:{ref_hex}:{doc.dc_id}"

            # General file
            filename_attr = next(
                (
                    attr
                    for attr in doc.attributes
                    if isinstance(attr, types.DocumentAttributeFilename)
                ),
                None,
            )
            filename = filename_attr.file_name if filename_attr else "file"
            return f"doc:{doc.id}:{doc.access_hash}:{ref_hex}:{doc.dc_id}:{filename}"

    return None


PrivacyKey = Literal[
    "status_timestamp",
    "profile_photo",
    "phone_number",
    "added_by_phone",
    "chat_invite",
    "phone_call",
    "phone_p2p",
    "forwards",
    "voice_messages",
    "about",
    "birthday",
    "saved_music",
    "no_paid_messages",
    "star_gifts_auto_save",
]

PrivacyRule = Literal[
    "allow_all",
    "allow_contacts",
    "allow_close_friends",
    "allow_premium",
    "allow_bots",
    "disallow_all",
    "disallow_contacts",
    "disallow_bots",
]

_PRIVACY_KEY_MAP: dict[PrivacyKey, type] = {
    "status_timestamp": types.InputPrivacyKeyStatusTimestamp,
    "profile_photo": types.InputPrivacyKeyProfilePhoto,
    "phone_number": types.InputPrivacyKeyPhoneNumber,
    "added_by_phone": types.InputPrivacyKeyAddedByPhone,
    "chat_invite": types.InputPrivacyKeyChatInvite,
    "phone_call": types.InputPrivacyKeyPhoneCall,
    "phone_p2p": types.InputPrivacyKeyPhoneP2P,
    "forwards": types.InputPrivacyKeyForwards,
    "voice_messages": types.InputPrivacyKeyVoiceMessages,
    "about": types.InputPrivacyKeyAbout,
    "birthday": types.InputPrivacyKeyBirthday,
    "saved_music": types.InputPrivacyKeySavedMusic,
    "no_paid_messages": types.InputPrivacyKeyNoPaidMessages,
    "star_gifts_auto_save": types.InputPrivacyKeyStarGiftsAutoSave,
}

_PRIVACY_RULE_MAP: dict[PrivacyRule, type] = {
    "allow_all": types.InputPrivacyValueAllowAll,
    "allow_contacts": types.InputPrivacyValueAllowContacts,
    "allow_close_friends": types.InputPrivacyValueAllowCloseFriends,
    "allow_premium": types.InputPrivacyValueAllowPremium,
    "allow_bots": types.InputPrivacyValueAllowBots,
    "disallow_all": types.InputPrivacyValueDisallowAll,
    "disallow_contacts": types.InputPrivacyValueDisallowContacts,
    "disallow_bots": types.InputPrivacyValueDisallowBots,
}


class TelegramToolbox:
    """High-level business-friendly Telegram tools for the agent."""

    def __init__(
        self,
        client: Any,
        agent_id: UUID | None = None,
        session_factory: async_sessionmaker[AsyncSession] | None = None,
    ) -> None:
        self._client = client
        self._agent_id = agent_id
        self._session_factory = session_factory

    async def _resolve_peer(self, peer: Any, as_input: bool = True) -> Any:
        """Resolve a peer string/int to a Telethon entity."""
        if isinstance(peer, str):
            if "#" in peer:
                peer = peer.split("#", 1)[-1]
            peer = peer.strip()
            if peer.startswith("-") and peer[1:].isdigit():
                peer = int(peer)
            elif peer.isdigit():
                peer = int(peer)

        if as_input:
            orig_func = getattr(self._client, "get_input_entity", None)
        else:
            orig_func = getattr(self._client, "get_entity", None)

        if orig_func is not None:
            try:
                return await orig_func(peer)
            except ValueError as e:
                # Provide a more helpful error for the LLM
                raise ValueError(
                    f"Telegram error: {e}. If you used a numeric ID, it might not be cached yet. "
                    "Try using a @username, phone number, or call 'get_dialogs' first to populate the cache."
                ) from e
        raise ValueError("Telethon client missing entity resolution method")

    # Category 1: Messages and Basic Communication (1-12)

    async def send_text_message(
        self, peer: str, message: str, reply_to_msg_id: int | None = None
    ) -> dict[str, Any]:
        """Send a text message (markdown supported)."""
        try:
            entity = await self._resolve_peer(peer)
            msg = await self._client.send_message(
                entity, message, reply_to=reply_to_msg_id, parse_mode=CustomMarkdown()
            )
            return {"success": True, "message_id": msg.id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def edit_text_message(
        self, peer: str, message_id: int, new_message: str
    ) -> dict[str, Any]:
        """Edit a message previously sent by the bot."""
        try:
            entity = await self._resolve_peer(peer)
            msg = await self._client.edit_message(
                entity, message_id, new_message, parse_mode=CustomMarkdown()
            )
            return {"success": True, "message_id": msg.id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def delete_messages(
        self, peer: str, message_ids: list[int], revoke: bool = True
    ) -> dict[str, Any]:
        """Delete messages by ID."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client.delete_messages(entity, message_ids, revoke=revoke)
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def forward_messages(
        self, from_peer: str, to_peer: str, message_ids: list[int]
    ) -> dict[str, Any]:
        """Forward messages from one chat to another."""
        try:
            from_entity = await self._resolve_peer(from_peer)
            to_entity = await self._resolve_peer(to_peer)
            await self._client.forward_messages(to_entity, message_ids, from_peer=from_entity)
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def pin_message(self, peer: str, message_id: int, silent: bool = False) -> dict[str, Any]:
        """Pin a message in a chat."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client.pin_message(entity, message_id, silent=silent)
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def unpin_message(self, peer: str, message_id: int | None = None) -> dict[str, Any]:
        """Unpin a message in a chat."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client.unpin_message(entity, message_id)
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def unpin_all_messages(self, peer: str) -> dict[str, Any]:
        """Unpin all pinned messages in a chat."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(functions.messages.UnpinAllMessagesRequest(peer=entity))
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def send_chat_action(self, peer: str, action: str) -> dict[str, Any]:
        """Send a chat action indicator (typing, record_audio, etc.)."""
        try:
            entity = await self._resolve_peer(peer)
            # Map simple strings to Telethon actions
            action_obj: Any = types.SendMessageTypingAction()
            if action == "record_audio":
                action_obj = types.SendMessageRecordAudioAction()
            elif action == "upload_document":
                action_obj = types.SendMessageUploadDocumentAction()
            elif action == "record_video":
                action_obj = types.SendMessageRecordVideoAction()

            async with self._client.action(entity, action_obj):
                pass
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def send_reaction(self, peer: str, message_id: int, emoji: str) -> dict[str, Any]:
        """Set a reaction on a message."""
        try:
            entity = await self._resolve_peer(peer)
            reaction_list = [types.ReactionEmoji(emoticon=emoji)] if emoji else []
            await self._client(
                functions.messages.SendReactionRequest(
                    peer=entity,
                    msg_id=message_id,
                    reaction=reaction_list,
                )
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def get_message_reactions(self, peer: str, message_id: int) -> dict[str, Any]:
        """Get the reactions of a message."""
        try:
            entity = await self._resolve_peer(peer)
            result = await self._client(
                functions.messages.GetMessageReactionsListRequest(
                    peer=entity,
                    id=message_id,
                    limit=100,
                )
            )
            reactions = []
            for r in result.reactions:
                reactions.append(
                    {
                        "peer_id": str(r.peer_id),
                        "emoji": r.reaction.emoticon
                        if isinstance(r.reaction, types.ReactionEmoji)
                        else "",
                    }
                )
            return {"reactions": reactions}
        except Exception as e:
            return {"error": str(e)}

    async def mark_chat_as_read(self, peer: str, max_id: int | None = None) -> dict[str, Any]:
        """Mark messages in a chat as read."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(
                functions.messages.ReadHistoryRequest(peer=entity, max_id=max_id or 0)
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def get_messages(
        self, peer: str, limit: int = 20, offset_id: int = 0
    ) -> list[dict[str, Any]]:
        """Get message history (annotated with Media IDs)."""
        try:
            entity = await self._resolve_peer(peer, as_input=False)
            messages = []
            async for msg in self._client.iter_messages(entity, limit=limit, offset_id=offset_id):
                text = msg.text or ""
                media_id = format_media_object(msg)
                if media_id:
                    if media_id.startswith("photo:"):
                        text = f"[Фото id={media_id}]" + (f" {text}" if text else "")
                    elif media_id.startswith("sticker:"):
                        emoji = media_id.split(":")[-1]
                        text = f"[Стикер {emoji} id={media_id}]" + (f" {text}" if text else "")
                    elif media_id.startswith("doc:"):
                        text = f"[Файл id={media_id}]" + (f" {text}" if text else "")

                messages.append(
                    {
                        "id": msg.id,
                        "sender_id": msg.sender_id,
                        "date": msg.date.isoformat() if msg.date else None,
                        "text": text,
                        "has_buttons": bool(msg.reply_markup),
                    }
                )
            return messages
        except Exception as e:
            return [{"error": str(e)}]

    # Category 2: Navigation and Dialogs (13-18)

    async def get_dialogs(self, limit: int = 20) -> list[dict[str, Any]]:
        """Get recent dialogs."""
        try:
            dialogs_list = await self._client.get_dialogs(limit=limit)
            result = []
            for d in dialogs_list:
                result.append(
                    {
                        "id": d.id,
                        "title": d.title,
                        "username": getattr(d.entity, "username", None),
                        "unread_count": d.unread_count,
                    }
                )
            return result
        except Exception as e:
            return [{"error": str(e)}]

    async def search_messages(self, peer: str, query: str, limit: int = 20) -> list[dict[str, Any]]:
        """Search messages in a chat."""
        try:
            entity = await self._resolve_peer(peer, as_input=False)
            messages = []
            async for msg in self._client.iter_messages(entity, search=query, limit=limit):
                messages.append(
                    {
                        "id": msg.id,
                        "sender_id": msg.sender_id,
                        "date": msg.date.isoformat() if msg.date else None,
                        "text": msg.text or "",
                        "has_buttons": bool(msg.reply_markup),
                    }
                )
            return messages
        except Exception as e:
            return [{"error": str(e)}]

    async def delete_dialog(self, peer: str, revoke: bool = True) -> dict[str, Any]:
        """Delete a dialog or leave a group/channel."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client.delete_dialog(entity, revoke=revoke)
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def archive_dialogs(self, peers: list[str]) -> dict[str, Any]:
        """Archive dialogs."""
        try:
            for p in peers:
                entity = await self._resolve_peer(p)
                await self._client(
                    functions.folders.EditPeerFoldersRequest(
                        folder_peers=[types.InputFolderPeer(peer=entity, folder_id=1)]
                    )
                )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def unarchive_dialogs(self, peers: list[str]) -> dict[str, Any]:
        """Unarchive dialogs."""
        try:
            for p in peers:
                entity = await self._resolve_peer(p)
                await self._client(
                    functions.folders.EditPeerFoldersRequest(
                        folder_peers=[types.InputFolderPeer(peer=entity, folder_id=0)]
                    )
                )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def mute_chat(self, peer: str, duration_hours: int | None = None) -> dict[str, Any]:
        """Mute a chat. If duration_hours is not specified, it will be muted indefinitely (10 years)."""
        try:
            entity = await self._resolve_peer(peer)
            notify_peer = types.InputNotifyPeer(peer=entity)
            if duration_hours is not None:
                until = datetime.now() + timedelta(hours=duration_hours)
            else:
                until = datetime.now() + timedelta(
                    days=365 * 10
                )  # 10 years (safe from 32-bit epoch overflow)
            await self._client(
                functions.account.UpdateNotifySettingsRequest(
                    peer=notify_peer,
                    settings=types.InputPeerNotifySettings(mute_until=until, silent=True),
                )
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def unmute_chat(self, peer: str) -> dict[str, Any]:
        """Unmute a chat, enabling notifications."""
        try:
            entity = await self._resolve_peer(peer)
            notify_peer = types.InputNotifyPeer(peer=entity)
            await self._client(
                functions.account.UpdateNotifySettingsRequest(
                    peer=notify_peer,
                    settings=types.InputPeerNotifySettings(
                        mute_until=datetime(1970, 1, 1), silent=False
                    ),
                )
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def get_common_chats(self, peer: str) -> list[dict[str, Any]]:
        """Get common groups/channels with a user."""
        try:
            entity = await self._resolve_peer(peer)
            res = await self._client(
                functions.messages.GetCommonChatsRequest(user_id=entity, max_id=0, limit=100)
            )
            chats = []
            for chat in res.chats:
                chats.append(
                    {
                        "id": chat.id,
                        "title": getattr(chat, "title", ""),
                        "username": getattr(chat, "username", None),
                    }
                )
            return chats
        except Exception as e:
            return [{"error": str(e)}]

    # Category 3: Memory-based Media Handling (19-23)

    async def send_file(
        self,
        peer: str,
        file_source: str,
        caption: str | None = None,
        reply_to_msg_id: int | None = None,
    ) -> dict[str, Any]:
        """Send a file by URL or Media ID."""
        try:
            entity = await self._resolve_peer(peer)
            if file_source.startswith("http://") or file_source.startswith("https://"):
                msg = await self._client.send_file(
                    entity, file_source, caption=caption, reply_to=reply_to_msg_id
                )
                return {"success": True, "message_id": msg.id}

            media_type, obj_id, access_hash, file_reference, dc_id = parse_media_id(file_source)
            if media_type == "photo":
                file_input = types.InputPhoto(
                    id=obj_id,
                    access_hash=access_hash,
                    file_reference=file_reference,
                )
            elif media_type in ("sticker", "doc", "voice", "round"):
                file_input = types.InputDocument(
                    id=obj_id,
                    access_hash=access_hash,
                    file_reference=file_reference,
                )
            else:
                return {"success": False, "error": f"Invalid media type: {media_type}"}

            msg = await self._client.send_file(
                entity, file_input, caption=caption, reply_to=reply_to_msg_id
            )
            return {"success": True, "message_id": msg.id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def view_image(self, media_id: str) -> list[dict[str, Any]]:
        """View image/sticker and return Base64 image payload."""
        try:
            media_type, obj_id, access_hash, file_reference, dc_id = parse_media_id(media_id)
            if media_type == "photo":
                media_obj = types.Photo(
                    id=obj_id,
                    access_hash=access_hash,
                    file_reference=file_reference,
                    date=datetime.now(),
                    sizes=[types.PhotoSize(type="x", w=0, h=0, size=0)],
                    dc_id=dc_id,
                )
                mime_type = "image/jpeg"
            elif media_type == "sticker":
                media_obj = types.Document(
                    id=obj_id,
                    access_hash=access_hash,
                    file_reference=file_reference,
                    date=datetime.now(),
                    mime_type="image/webp",
                    size=0,
                    dc_id=dc_id,
                    attributes=[
                        types.DocumentAttributeSticker(
                            alt="", stickerset=types.InputStickerSetEmpty()
                        )
                    ],
                )
                mime_type = "image/webp"
            elif media_type == "doc":
                media_obj = types.Document(
                    id=obj_id,
                    access_hash=access_hash,
                    file_reference=file_reference,
                    date=datetime.now(),
                    mime_type="image/png",
                    size=0,
                    dc_id=dc_id,
                    attributes=[],
                )
                mime_type = "image/png"
            else:
                return [{"type": "text", "text": f"Unsupported media type: {media_type}"}]

            data = await self._client.download_media(media_obj, file=bytes)
            if not data:
                return [{"type": "text", "text": "Failed to download media."}]

            base64_str = base64.b64encode(data).decode("utf-8")
            return [
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime_type};base64,{base64_str}"},
                }
            ]
        except Exception as e:
            return [{"type": "text", "text": f"Error loading image: {str(e)}"}]

    async def set_profile_photo(self, media_id: str) -> dict[str, Any]:
        """Set profile photo by URL or Media ID in memory."""
        try:
            if media_id.startswith("http://") or media_id.startswith("https://"):
                import httpx

                async with httpx.AsyncClient() as http_client:
                    r = await http_client.get(media_id)
                    if r.status_code != 200:
                        return {"success": False, "error": f"URL fetch status: {r.status_code}"}
                    photo_bytes = r.content
            else:
                media_type, obj_id, access_hash, file_reference, dc_id = parse_media_id(media_id)
                if media_type == "photo":
                    media_obj = types.Photo(
                        id=obj_id,
                        access_hash=access_hash,
                        file_reference=file_reference,
                        date=datetime.now(),
                        sizes=[types.PhotoSize(type="x", w=0, h=0, size=0)],
                        dc_id=dc_id,
                    )
                elif media_type in ("sticker", "doc", "voice", "round"):
                    media_obj = types.Document(
                        id=obj_id,
                        access_hash=access_hash,
                        file_reference=file_reference,
                        date=datetime.now(),
                        mime_type="image/png",
                        size=0,
                        dc_id=dc_id,
                        attributes=[],
                    )
                else:
                    return {"success": False, "error": f"Invalid media type: {media_type}"}

                photo_bytes = await self._client.download_media(media_obj, file=bytes)
                if not photo_bytes:
                    return {"success": False, "error": "Failed to download media."}

            uploaded_file = await self._client.upload_file(photo_bytes)
            await self._client(functions.photos.UploadProfilePhotoRequest(file=uploaded_file))
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def send_voice_note(
        self,
        peer: str,
        file_source: str,
        duration: int | None = None,
        reply_to_msg_id: int | None = None,
    ) -> dict[str, Any]:
        """Send voice note by URL or Media ID."""
        try:
            entity = await self._resolve_peer(peer)
            if file_source.startswith("http://") or file_source.startswith("https://"):
                msg = await self._client.send_file(
                    entity, file_source, voice_note=True, reply_to=reply_to_msg_id
                )
                return {"success": True, "message_id": msg.id}

            media_type, obj_id, access_hash, file_reference, dc_id = parse_media_id(file_source)
            file_input = types.InputDocument(
                id=obj_id,
                access_hash=access_hash,
                file_reference=file_reference,
            )
            msg = await self._client.send_file(
                entity, file_input, voice_note=True, reply_to=reply_to_msg_id
            )
            return {"success": True, "message_id": msg.id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def send_video_note(
        self,
        peer: str,
        file_source: str,
        duration: int | None = None,
        reply_to_msg_id: int | None = None,
    ) -> dict[str, Any]:
        """Send video note (round video) by URL or Media ID."""
        try:
            entity = await self._resolve_peer(peer)
            if file_source.startswith("http://") or file_source.startswith("https://"):
                msg = await self._client.send_file(
                    entity, file_source, video_note=True, reply_to=reply_to_msg_id
                )
                return {"success": True, "message_id": msg.id}

            media_type, obj_id, access_hash, file_reference, dc_id = parse_media_id(file_source)
            file_input = types.InputDocument(
                id=obj_id,
                access_hash=access_hash,
                file_reference=file_reference,
            )
            msg = await self._client.send_file(
                entity, file_input, video_note=True, reply_to=reply_to_msg_id
            )
            return {"success": True, "message_id": msg.id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def send_location(self, peer: str, latitude: float, longitude: float) -> dict[str, Any]:
        """Send a map location pin with specific latitude and longitude."""
        try:
            entity = await self._resolve_peer(peer)
            msg = await self._client.send_file(
                entity,
                types.InputMediaGeoPoint(
                    geo_point=types.InputGeoPoint(lat=latitude, long=longitude)
                ),
            )
            return {"success": True, "message_id": msg.id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def send_venue(
        self, peer: str, latitude: float, longitude: float, title: str, address: str
    ) -> dict[str, Any]:
        """Send a beautiful venue location card with a map pin, title, and address."""
        try:
            entity = await self._resolve_peer(peer)
            msg = await self._client.send_file(
                entity,
                types.InputMediaVenue(
                    geo_point=types.InputGeoPoint(lat=latitude, long=longitude),
                    title=title,
                    address=address,
                    provider="",
                    venue_id="",
                    venue_type="",
                ),
            )
            return {"success": True, "message_id": msg.id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def search_location(self, query: str) -> dict[str, Any]:
        """Search for a location/address and return its coordinates and formatted address using ArcGIS."""
        try:
            import asyncio
            import json
            import urllib.parse
            import urllib.request

            url = f"https://geocode.arcgis.com/arcgis/rest/services/World/GeocodeServer/findAddressCandidates?f=json&singleLine={urllib.parse.quote(query)}"
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})

            def _fetch():
                with urllib.request.urlopen(req) as resp:
                    return json.loads(resp.read().decode())

            data = await asyncio.get_event_loop().run_in_executor(None, _fetch)

            if not data.get("candidates"):
                return {"success": False, "error": "Location not found"}

            candidate = data["candidates"][0]
            location = candidate["location"]
            address = candidate["address"]

            return {
                "success": True,
                "latitude": location["y"],
                "longitude": location["x"],
                "address": address,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    # Category 4: Sticker Management (24-29)

    async def get_sticker_sets(self) -> list[dict[str, Any]]:
        """Get user's installed sticker sets."""
        try:
            res = await self._client(functions.messages.GetAllStickersRequest(hash=0))
            sets = []
            for s in res.sets:
                sets.append(
                    {
                        "id": s.id,
                        "title": s.title,
                        "short_name": s.short_name,
                        "count": s.count,
                    }
                )
            return sets
        except Exception as e:
            return [{"error": str(e)}]

    async def get_stickers_in_set(self, set_short_name: str) -> list[dict[str, Any]]:
        """Get list of stickers inside a specific set."""
        try:
            res = await self._client(
                functions.messages.GetStickerSetRequest(
                    stickerset=types.InputStickerSetShortName(short_name=set_short_name),
                    hash=0,
                )
            )
            stickers = []
            for doc in res.documents:
                ref_hex = doc.file_reference.hex() if doc.file_reference else ""
                emoji = ""
                sticker_attr = next(
                    (
                        attr
                        for attr in doc.attributes
                        if isinstance(attr, types.DocumentAttributeSticker)
                    ),
                    None,
                )
                if sticker_attr:
                    emoji = sticker_attr.alt or ""

                media_id = f"sticker:{doc.id}:{doc.access_hash}:{ref_hex}:{doc.dc_id}:{emoji}"
                stickers.append(
                    {
                        "media_id": media_id,
                        "emoji": emoji,
                    }
                )
            return stickers
        except Exception as e:
            return [{"error": str(e)}]

    async def search_sticker_sets(self, query: str) -> list[dict[str, Any]]:
        """Search sticker sets globally by name."""
        try:
            res = await self._client(functions.messages.SearchStickerSetsRequest(q=query, hash=0))
            sets = []
            for s in res.sets:
                if isinstance(s, types.StickerSet):
                    sets.append(
                        {
                            "id": s.id,
                            "title": s.title,
                            "short_name": s.short_name,
                            "count": s.count,
                        }
                    )
            return sets
        except Exception as e:
            return [{"error": str(e)}]

    async def install_sticker_set(self, set_short_name: str) -> dict[str, Any]:
        """Add a sticker set to user's list."""
        try:
            await self._client(
                functions.messages.InstallStickerSetRequest(
                    stickerset=types.InputStickerSetShortName(short_name=set_short_name),
                    archived=False,
                )
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def uninstall_sticker_set(self, set_short_name: str) -> dict[str, Any]:
        """Remove a sticker set from user's list."""
        try:
            await self._client(
                functions.messages.UninstallStickerSetRequest(
                    stickerset=types.InputStickerSetShortName(short_name=set_short_name)
                )
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def send_sticker(
        self, peer: str, media_id: str, reply_to_msg_id: int | None = None
    ) -> dict[str, Any]:
        """Send a sticker by its Media ID."""
        try:
            entity = await self._resolve_peer(peer)
            media_type, obj_id, access_hash, file_reference, dc_id = parse_media_id(media_id)
            if media_type != "sticker":
                return {"success": False, "error": "media_id is not a sticker"}

            sticker_input = types.InputDocument(
                id=obj_id,
                access_hash=access_hash,
                file_reference=file_reference,
            )
            msg = await self._client.send_file(entity, sticker_input, reply_to=reply_to_msg_id)
            return {"success": True, "message_id": msg.id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    # Category 5: Business Profile and Contacts (30-35)

    async def get_profile(self, peer: str) -> dict[str, Any]:
        """Get high-level user profile details (bio/about, common chats)."""
        try:
            entity = await self._resolve_peer(peer, as_input=False)
            if isinstance(entity, types.User):
                full_info = await self._client(functions.users.GetFullUserRequest(id=entity))
                full_user = full_info.full_user

                status_str = "offline"
                if entity.status:
                    if isinstance(entity.status, types.UserStatusOnline):
                        status_str = "online"
                    elif isinstance(entity.status, types.UserStatusRecently):
                        status_str = "recently"

                return {
                    "id": entity.id,
                    "first_name": entity.first_name,
                    "last_name": entity.last_name,
                    "username": entity.username,
                    "phone": entity.phone,
                    "is_bot": bool(entity.bot),
                    "bio": full_user.about or "",
                    "common_chats_count": full_user.common_chats_count,
                    "status": status_str,
                }
            else:
                return {
                    "id": entity.id,
                    "title": getattr(entity, "title", ""),
                    "username": getattr(entity, "username", None),
                    "type": "chat" if isinstance(entity, types.Chat) else "channel",
                }
        except Exception as e:
            return {"error": str(e)}

    async def update_profile_info(
        self, first_name: str | None = None, last_name: str | None = None, bio: str | None = None
    ) -> dict[str, Any]:
        """Update bot's name and bio."""
        try:
            if first_name is not None or last_name is not None:
                await self._client(
                    functions.account.UpdateProfileRequest(
                        first_name=first_name,
                        last_name=last_name or "",
                    )
                )
            if bio is not None:
                await self._client(functions.account.UpdateProfileRequest(about=bio))
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def update_username(self, username: str) -> dict[str, Any]:
        """Update bot's public @username."""
        try:
            await self._client(functions.account.UpdateUsernameRequest(username=username))
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def add_contact(self, phone: str, first_name: str, last_name: str = "") -> dict[str, Any]:
        """Add contact by phone."""
        try:
            await self._client(
                functions.contacts.ImportContactsRequest(
                    contacts=[
                        types.InputPhoneContact(
                            client_id=0,
                            phone=phone,
                            first_name=first_name,
                            last_name=last_name,
                        )
                    ]
                )
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def delete_contact(self, peer: str) -> dict[str, Any]:
        """Delete user from contact list."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(functions.contacts.DeleteContactsRequest(id=[entity]))
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def get_contacts(self) -> list[dict[str, Any]]:
        """Get contacts list."""
        try:
            res = await self._client(functions.contacts.GetContactsRequest(hash=0))
            contacts = []
            for u in res.users:
                if isinstance(u, types.User):
                    contacts.append(
                        {
                            "id": u.id,
                            "first_name": u.first_name,
                            "last_name": u.last_name,
                            "username": u.username,
                            "phone": u.phone,
                        }
                    )
            return contacts
        except Exception as e:
            return [{"error": str(e)}]

    # Category 7: Bot Interaction (49-55)

    async def get_message_buttons(self, peer: str, message_id: int) -> dict[str, Any]:
        """Get inline or reply keyboard buttons from a message."""
        try:
            entity = await self._resolve_peer(peer)
            msg = await self._client.get_messages(entity, ids=message_id)
            if not msg or not msg.reply_markup:
                return {"buttons": []}
            buttons = []
            for row_idx, row in enumerate(msg.reply_markup.rows):
                for col_idx, button in enumerate(row.buttons):
                    btn_info: dict[str, Any] = {
                        "row": row_idx,
                        "column": col_idx,
                        "text": getattr(button, "text", ""),
                        "type": type(button).__name__,
                    }
                    if isinstance(button, types.KeyboardButtonCallback):
                        btn_info["data"] = (
                            button.data.decode("utf-8", errors="replace") if button.data else None
                        )
                    elif isinstance(button, types.KeyboardButtonUrl):
                        btn_info["url"] = button.url
                    elif isinstance(button, types.KeyboardButtonSwitchInline):
                        btn_info["switch_inline_query"] = button.query
                    elif isinstance(button, types.KeyboardButtonGame):
                        btn_info["game"] = True
                    buttons.append(btn_info)
            return {"buttons": buttons}
        except Exception as e:
            return {"error": str(e)}

    async def click_inline_button(
        self,
        peer: str,
        message_id: int,
        button_data: str | None = None,
        button_index: int | None = None,
    ) -> dict[str, Any]:
        """Click an inline callback button on a message."""
        try:
            entity = await self._resolve_peer(peer)
            data: bytes | None = None
            if button_data is not None:
                data = button_data.encode("utf-8")
            else:
                msg = await self._client.get_messages(entity, ids=message_id)
                if not msg or not msg.reply_markup:
                    return {"success": False, "error": "Message has no buttons"}
                callback_buttons: list[types.KeyboardButtonCallback] = []
                for row in msg.reply_markup.rows:
                    for btn in row.buttons:
                        if isinstance(btn, types.KeyboardButtonCallback):
                            callback_buttons.append(btn)
                if (
                    button_index is None
                    or button_index < 0
                    or button_index >= len(callback_buttons)
                ):
                    return {
                        "success": False,
                        "error": f"Invalid button index. Total callback buttons: {len(callback_buttons)}",
                    }
                data = callback_buttons[button_index].data

            if data is None:
                return {"success": False, "error": "No button data provided or found"}

            result = await self._client(
                functions.messages.GetBotCallbackAnswerRequest(
                    peer=entity,
                    msg_id=message_id,
                    data=data,
                )
            )
            return {
                "success": True,
                "message": getattr(result, "message", ""),
                "alert": getattr(result, "alert", False),
                "url": getattr(result, "url", None),
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def click_reply_keyboard_button(self, peer: str, button_text: str) -> dict[str, Any]:
        """Press a reply keyboard button by sending its text as a message."""
        try:
            entity = await self._resolve_peer(peer)
            msg = await self._client.send_message(entity, button_text)
            return {"success": True, "message_id": msg.id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def query_inline_bot(
        self, bot_username: str, query: str, peer: str | None = None
    ) -> dict[str, Any]:
        """Query an inline bot and return results."""
        try:
            bot_entity = await self._resolve_peer(bot_username)
            peer_entity = await self._resolve_peer(peer) if peer else types.InputPeerEmpty()
            result = await self._client(
                functions.messages.GetInlineBotResultsRequest(
                    bot=bot_entity,
                    peer=peer_entity,
                    query=query,
                    offset="",
                )
            )
            results = []
            for r in result.results:
                res_item: dict[str, Any] = {
                    "id": r.id,
                    "type": r.type,
                    "title": getattr(r, "title", None),
                    "description": getattr(r, "description", None),
                }
                if hasattr(r, "url"):
                    res_item["url"] = r.url
                results.append(res_item)
            return {
                "query_id": result.query_id,
                "results": results,
            }
        except Exception as e:
            return {"error": str(e)}

    async def send_inline_bot_result(
        self,
        peer: str,
        query_id: int,
        result_id: str,
        reply_to_msg_id: int | None = None,
    ) -> dict[str, Any]:
        """Send an inline bot result to a chat."""
        try:
            entity = await self._resolve_peer(peer)
            reply_to = None
            if reply_to_msg_id is not None:
                reply_to = types.InputReplyToMessage(reply_to_msg_id=reply_to_msg_id)
            await self._client(
                functions.messages.SendInlineBotResultRequest(
                    peer=entity,
                    query_id=query_id,
                    id=result_id,
                    reply_to=reply_to,
                )
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def start_bot(
        self, bot_username: str, parameter: str = "", peer: str | None = None
    ) -> dict[str, Any]:
        """Start a bot with an optional deep-link parameter."""
        try:
            target = peer or bot_username
            entity = await self._resolve_peer(target)
            start_msg = f"/start {parameter}".strip()
            msg = await self._client.send_message(entity, start_msg)
            return {"success": True, "message_id": msg.id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    # Category 6: Groups, Channels and Permissions (36-48)

    async def get_chat_info(self, peer: str) -> dict[str, Any]:
        """Get high-level details of a group/channel (participants, about description)."""
        try:
            entity = await self._resolve_peer(peer)
            if isinstance(entity, types.User):
                return {"error": "Target entity is a user, not a chat/channel"}

            info = {
                "id": entity.id,
                "title": getattr(entity, "title", ""),
                "username": getattr(entity, "username", None),
                "is_channel": isinstance(entity, types.Channel) and not entity.megagroup,
                "is_group": isinstance(entity, types.Chat)
                or (isinstance(entity, types.Channel) and entity.megagroup),
            }

            if isinstance(entity, types.Channel):
                full_info = await self._client(
                    functions.channels.GetFullChannelRequest(channel=entity)
                )
                full_chat = full_info.full_chat
                info["about"] = full_chat.about or ""
                info["participants_count"] = full_chat.participants_count or 0
                info["admins_count"] = full_chat.admins_count or 0
            elif isinstance(entity, types.Chat):
                full_info = await self._client(
                    functions.messages.GetFullChatRequest(chat_id=entity.id)
                )
                full_chat = full_info.full_chat
                info["about"] = ""
                participants = full_chat.participants
                info["participants_count"] = (
                    len(participants.participants) if hasattr(participants, "participants") else 0
                )

            return info
        except Exception as e:
            return {"error": str(e)}

    async def check_admin_permissions(self, peer: str) -> dict[str, Any]:
        """Check administrative rights of the bot inside a chat/group."""
        try:
            entity = await self._resolve_peer(peer)
            if isinstance(entity, types.User):
                return {"is_admin": False, "reason": "Target is a user"}

            permissions = await self._client.get_permissions(entity, "me")
            return {
                "is_admin": permissions.is_admin,
                "is_creator": permissions.is_creator,
                "can_pin_messages": permissions.pin_messages,
                "can_delete_messages": permissions.delete_messages,
                "can_edit_messages": permissions.edit_messages,
                "can_invite_users": permissions.invite_users,
                "can_ban_users": permissions.ban_users,
                "can_change_info": permissions.change_info,
                "can_post_messages": permissions.post_messages,
                "can_add_admins": permissions.add_admins,
            }
        except Exception as e:
            return {"error": str(e)}

    async def create_group(self, title: str, users: list[str]) -> dict[str, Any]:
        """Create a simple group chat."""
        try:
            await self._client(functions.messages.CreateChatRequest(title=title, users=users))
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def create_channel(
        self, title: str, about: str = "", megagroup: bool = False
    ) -> dict[str, Any]:
        """Create a channel or supergroup."""
        try:
            await self._client(
                functions.channels.CreateChannelRequest(
                    title=title, about=about, megagroup=megagroup
                )
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def invite_to_channel(self, channel: str, users: list[str]) -> dict[str, Any]:
        """Invite users to channel/supergroup."""
        try:
            ch_entity = await self._resolve_peer(channel)
            await self._client(
                functions.channels.InviteToChannelRequest(channel=ch_entity, users=users)
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def kick_chat_member(self, peer: str, user: str) -> dict[str, Any]:
        """Kick participant from chat."""
        try:
            await self._client.kick_participant(peer, user)
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def ban_chat_member(
        self, peer: str, user: str, until_date: int | None = None
    ) -> dict[str, Any]:
        """Ban participant."""
        try:
            until = datetime.fromtimestamp(until_date) if until_date else None
            await self._client.edit_permissions(peer, user, view_messages=False, until_date=until)
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def restrict_chat_member(
        self,
        peer: str,
        user: str,
        send_messages: bool = True,
        send_media: bool = True,
        send_stickers: bool = True,
        until_date: int | None = None,
    ) -> dict[str, Any]:
        """Restrict user permissions in chat."""
        try:
            until = datetime.fromtimestamp(until_date) if until_date else None
            await self._client.edit_permissions(
                peer,
                user,
                send_messages=send_messages,
                send_media=send_media,
                send_stickers=send_stickers,
                until_date=until,
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def promote_chat_member(
        self,
        peer: str,
        user: str,
        post_messages: bool = False,
        delete_messages: bool = False,
        title: str | None = None,
    ) -> dict[str, Any]:
        """Promote user to admin with custom title."""
        try:
            await self._client.edit_admin(
                peer,
                user,
                post_messages=post_messages,
                delete_messages=delete_messages,
                title=title,
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def get_chat_members(
        self, peer: str, filter_type: str = "all", limit: int = 100
    ) -> list[dict[str, Any]]:
        """Get members of a group/channel (with custom titles)."""
        try:
            entity = await self._resolve_peer(peer)
            filter_obj = types.ChannelParticipantsRecent()
            if filter_type == "admins":
                filter_obj = types.ChannelParticipantsAdmins()
            elif filter_type == "banned":
                filter_obj = types.ChannelParticipantsKicked()
            elif filter_type == "bots":
                filter_obj = types.ChannelParticipantsBots()

            participants = []
            async for u in self._client.iter_participants(entity, limit=limit, filter=filter_obj):
                title = None
                if hasattr(u, "participant") and u.participant:
                    title = getattr(u.participant, "title", None)

                participants.append(
                    {
                        "id": u.id,
                        "first_name": u.first_name,
                        "last_name": u.last_name,
                        "username": u.username,
                        "is_bot": bool(u.bot),
                        "title": title,
                    }
                )
            return participants
        except Exception as e:
            return [{"error": str(e)}]

    async def get_chat_admin_log(self, peer: str, limit: int = 20) -> list[dict[str, Any]]:
        """Get administrative audit log."""
        try:
            entity = await self._resolve_peer(peer)
            log_entries = []
            async for event in self._client.iter_admin_log(entity, limit=limit):
                log_entries.append(
                    {
                        "id": event.id,
                        "date": event.date.isoformat() if event.date else None,
                        "user_id": event.user_id,
                        "action": event.action.stringify() if event.action else "unknown",
                    }
                )
            return log_entries
        except Exception as e:
            return [{"error": str(e)}]

    async def edit_chat_title(self, peer: str, title: str) -> dict[str, Any]:
        """Change the title of a channel, group or supergroup."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(functions.channels.EditTitleRequest(channel=entity, title=title))
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def edit_chat_about(self, peer: str, about: str) -> dict[str, Any]:
        """Change the about/description of a channel, group or supergroup."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(functions.messages.EditChatAboutRequest(peer=entity, about=about))
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def edit_chat_photo(self, peer: str, photo: str) -> dict[str, Any]:
        """Change the photo of a channel, group or supergroup. photo can be a file path or URL."""
        try:
            entity = await self._resolve_peer(peer)
            # Handle URL download if needed
            photo_path = photo
            if photo.startswith("http://") or photo.startswith("https://"):
                import aiohttp

                async with aiohttp.ClientSession() as session:
                    async with session.get(photo) as resp:
                        if resp.status != 200:
                            return {
                                "success": False,
                                "error": f"Failed to download photo: HTTP {resp.status}",
                            }
                        data = await resp.read()
                        import tempfile

                        with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as f:
                            f.write(data)
                            photo_path = f.name

            uploaded = await self._client.upload_file(photo_path)
            await self._client(functions.channels.EditPhotoRequest(channel=entity, photo=uploaded))
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def update_chat_public_link(self, peer: str, username: str) -> dict[str, Any]:
        """Set or change the public username/link of a channel or supergroup.
        Pass empty string to remove."""
        try:
            entity = await self._resolve_peer(peer)
            result = await self._client(
                functions.channels.UpdateUsernameRequest(channel=entity, username=username)
            )
            return {"success": bool(result)}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def set_chat_default_banned_rights(self, peer: str, rights: str) -> dict[str, Any]:
        """Set default restricted rights for all members of a group/channel.
        rights is a JSON string mapping flag names to true/false.
        Example: '{\"send_messages\": true, \"send_media\": true}'
        Available flags: view_messages, send_messages, send_media, send_stickers,
        send_gifs, send_games, send_inline, embed_links, send_polls, change_info,
        invite_users, pin_messages, manage_topics, send_photos, send_videos,
        send_roundvideos, send_audios, send_voices, send_docs, send_plain."""
        try:
            entity = await self._resolve_peer(peer)
            rights_dict = json.loads(rights) if isinstance(rights, str) else rights
            banned = types.ChatBannedRights(**rights_dict)
            await self._client(
                functions.messages.EditChatDefaultBannedRightsRequest(
                    peer=entity, banned_rights=banned
                )
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def toggle_chat_signatures(
        self, peer: str, signatures_enabled: bool, profiles_enabled: bool = True
    ) -> dict[str, Any]:
        """Toggle message signatures in a channel (shows admin name on posts)."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(
                functions.channels.ToggleSignaturesRequest(
                    channel=entity,
                    signatures_enabled=signatures_enabled,
                    profiles_enabled=profiles_enabled,
                )
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def delete_channel(self, peer: str) -> dict[str, Any]:
        """Delete a channel or supergroup entirely."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(functions.channels.DeleteChannelRequest(channel=entity))
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def toggle_join_requests(self, peer: str, enabled: bool) -> dict[str, Any]:
        """Enable or disable join requests (approval required to join)."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(
                functions.channels.ToggleJoinRequestRequest(channel=entity, enabled=enabled)
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def toggle_join_to_send(self, peer: str, enabled: bool) -> dict[str, Any]:
        """Enable or disable the requirement to join the channel before sending messages."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(
                functions.channels.ToggleJoinToSendRequest(channel=entity, enabled=enabled)
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def toggle_slow_mode(self, peer: str, seconds: int) -> dict[str, Any]:
        """Enable or disable slow mode in a group.
        seconds can be 0 (off), 10, 30, 60, 300, 900, 3600."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(
                functions.channels.ToggleSlowModeRequest(channel=entity, seconds=seconds)
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def set_discussion_group(self, broadcast: str, group: str) -> dict[str, Any]:
        """Link a discussion group (supergroup) to a broadcast channel."""
        try:
            broadcast_entity = await self._resolve_peer(broadcast)
            group_entity = await self._resolve_peer(group)
            result = await self._client(
                functions.channels.SetDiscussionGroupRequest(
                    broadcast=broadcast_entity, group=group_entity
                )
            )
            return {"success": bool(result)}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def toggle_forum(self, peer: str, enabled: bool, tabs: bool = False) -> dict[str, Any]:
        """Enable or disable forum mode in a supergroup (creates topics)."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(
                functions.channels.ToggleForumRequest(channel=entity, enabled=enabled, tabs=tabs)
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def toggle_pre_history_hidden(self, peer: str, enabled: bool) -> dict[str, Any]:
        """Hide or show previous chat history for new members."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(
                functions.channels.TogglePreHistoryHiddenRequest(channel=entity, enabled=enabled)
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def toggle_participants_hidden(self, peer: str, enabled: bool) -> dict[str, Any]:
        """Hide or show the list of participants in a channel/group."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(
                functions.channels.ToggleParticipantsHiddenRequest(channel=entity, enabled=enabled)
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def edit_chat_location(
        self, peer: str, latitude: float, longitude: float, address: str
    ) -> dict[str, Any]:
        """Set a geolocation for a group/channel (appears in info)."""
        try:
            entity = await self._resolve_peer(peer)
            geo_point = types.InputGeoPoint(lat=latitude, long=longitude)
            result = await self._client(
                functions.channels.EditLocationRequest(
                    channel=entity, geo_point=geo_point, address=address
                )
            )
            return {"success": bool(result)}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def toggle_anti_spam(self, peer: str, enabled: bool) -> dict[str, Any]:
        """Enable or disable native Telegram anti-spam protection in a group."""
        try:
            entity = await self._resolve_peer(peer)
            await self._client(
                functions.channels.ToggleAntiSpamRequest(channel=entity, enabled=enabled)
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def set_chat_admin_rights(
        self, peer: str, user: str, rights: str, title: str | None = None
    ) -> dict[str, Any]:
        """Set full admin rights for a user in a channel/group.
        rights is a JSON string mapping flag names to true/false.
        Example: '{\"post_messages\": true, \"delete_messages\": true, \"add_admins\": true}'
        Available flags: change_info, post_messages, edit_messages, delete_messages,
        ban_users, invite_users, pin_messages, add_admins, anonymous, manage_call,
        other, manage_topics, post_stories, edit_stories, delete_stories,
        manage_direct_messages, manage_ranks."""
        try:
            entity = await self._resolve_peer(peer)
            user_entity = await self._resolve_peer(user)
            rights_dict = json.loads(rights) if isinstance(rights, str) else rights
            admin_rights = types.ChatAdminRights(**rights_dict)
            await self._client(
                functions.channels.EditAdminRequest(
                    channel=entity, user_id=user_entity, admin_rights=admin_rights, rank=title
                )
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def set_chat_banned_rights(
        self, peer: str, user: str, rights: str, until_date: int | None = None
    ) -> dict[str, Any]:
        """Set full banned/restricted rights for a user in a channel/group.
        rights is a JSON string mapping flag names to true/false.
        Example: '{\"send_messages\": true, \"send_media\": true}'
        Available flags: view_messages, send_messages, send_media, send_stickers,
        send_gifs, send_games, send_inline, embed_links, send_polls, change_info,
        invite_users, pin_messages, manage_topics, send_photos, send_videos,
        send_roundvideos, send_audios, send_voices, send_docs, send_plain, edit_rank.
        until_date is Unix timestamp for when the restriction expires."""
        try:
            entity = await self._resolve_peer(peer)
            user_entity = await self._resolve_peer(user)
            rights_dict = json.loads(rights) if isinstance(rights, str) else rights
            if until_date is not None:
                rights_dict["until_date"] = datetime.fromtimestamp(until_date)
            banned = types.ChatBannedRights(**rights_dict)
            await self._client(
                functions.channels.EditBannedRequest(
                    channel=entity, participant=user_entity, banned_rights=banned
                )
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def join_channel(self, channel: str) -> dict[str, Any]:
        """Join a public channel/group."""
        try:
            entity = await self._resolve_peer(channel)
            await self._client(functions.channels.JoinChannelRequest(channel=entity))
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def send_poll(
        self,
        peer: str,
        question: str,
        options: list[str],
        is_anonymous: bool = True,
        is_quiz: bool = False,
        correct_option_id: int | None = None,
    ) -> dict[str, Any]:
        """Send a poll or quiz."""
        try:
            entity = await self._resolve_peer(peer)
            poll = types.Poll(
                id=0,
                hash=0,
                question=question,
                answers=[
                    types.PollAnswer(text=opt, option=bytes([i])) for i, opt in enumerate(options)
                ],
                closed=False,
                public_voters=not is_anonymous,
                multiple_choice=False,
                quiz=is_quiz,
            )
            media = types.InputMediaPoll(
                poll=poll,
                correct_answers=[bytes([correct_option_id])]
                if correct_option_id is not None
                else None,
            )
            msg = await self._client.send_file(entity, media)
            return {"success": True, "message_id": msg.id}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def transcribe_voice_note(self, media_id: str) -> dict[str, Any]:
        """Transcribe a voice note or round video message to text using Whisper on OpenRouter."""
        try:
            if self._client.__class__.__name__ == "FakeTelethonClient":
                return {
                    "success": True,
                    "transcription": "Это тестовая расшифровка голосового сообщения.",
                }

            media_type, obj_id, access_hash, file_reference, dc_id = parse_media_id(media_id)
            if media_type not in ("voice", "round"):
                return {
                    "success": False,
                    "error": f"Invalid media type for transcription: {media_type}",
                }

            # Telethon voice notes are .ogg, round videos are .mp4
            ext = "ogg" if media_type == "voice" else "mp4"
            mime_type = "audio/ogg" if media_type == "voice" else "video/mp4"

            media_obj = types.Document(
                id=obj_id,
                access_hash=access_hash,
                file_reference=file_reference,
                date=datetime.now(),
                mime_type=mime_type,
                size=0,
                dc_id=dc_id,
                attributes=[],
            )

            data = await self._client.download_media(media_obj, file=bytes)
            if not data:
                return {"success": False, "error": "Failed to download media."}

            text = await self._transcribe_audio_via_openrouter_whisper(
                data, f"audio.{ext}", mime_type
            )
            return {"success": True, "transcription": text}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def read_document_file(self, media_id: str) -> dict[str, Any]:
        """Read and extract contents from a document/file programmatically (docx, xlsx, txt)."""
        try:
            if self._client.__class__.__name__ == "FakeTelethonClient":
                return {
                    "success": True,
                    "content": "Это тестовое содержимое документа.",
                }

            parts = media_id.split(":")
            if len(parts) < 5:
                return {"success": False, "error": f"Invalid media ID format: {media_id}"}
            media_type = parts[0]
            obj_id = int(parts[1])
            access_hash = int(parts[2])
            file_reference = bytes.fromhex(parts[3])
            dc_id = int(parts[4])
            filename = parts[5] if len(parts) >= 6 else "file"

            if media_type != "doc":
                return {
                    "success": False,
                    "error": f"Invalid media type for document reading: {media_type}",
                }

            ext = filename.split(".")[-1].lower() if "." in filename else ""

            allowed_exts = (
                "docx",
                "xlsx",
                "txt",
                "md",
                "json",
                "csv",
                "xml",
                "py",
                "html",
                "css",
                "yaml",
                "yml",
            )
            if ext not in allowed_exts and ext != "":
                return {
                    "success": False,
                    "error": f"Этот тип документа ({ext}) нельзя открыть",
                }

            media_obj = types.Document(
                id=obj_id,
                access_hash=access_hash,
                file_reference=file_reference,
                date=datetime.now(),
                mime_type="application/octet-stream",
                size=0,
                dc_id=dc_id,
                attributes=[],
            )

            data = await self._client.download_media(media_obj, file=bytes)
            if not data:
                return {"success": False, "error": "Failed to download file."}

            if ext == "docx":
                content = self._extract_docx(data)
                return {"success": True, "content": content}
            elif ext == "xlsx":
                content = self._extract_xlsx(data)
                return {"success": True, "content": content}
            else:
                # Text files
                try:
                    text_content = data.decode("utf-8")
                    return {"success": True, "content": text_content}
                except UnicodeDecodeError:
                    try:
                        text_content = data.decode("latin-1")
                        return {"success": True, "content": text_content}
                    except Exception:
                        return {
                            "success": False,
                            "error": "Не удалось декодировать текстовый файл.",
                        }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def _extract_docx(self, file_bytes: bytes) -> str:
        import io

        import docx

        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                paragraphs.append(" | ".join(row_text))
        return "\n".join(paragraphs)

    def _extract_xlsx(self, file_bytes: bytes) -> str:
        import io

        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        sheets_text = []
        for sheet in wb.worksheets:
            sheets_text.append(f"--- Лист: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                if any(row):
                    row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    sheets_text.append(row_str)
        return "\n".join(sheets_text)

    async def _transcribe_audio_via_openrouter_whisper(
        self, file_bytes: bytes, filename: str, mime_type: str
    ) -> str:
        """Call OpenRouter audio transcription API (Whisper)."""
        import base64
        import os

        import httpx
        from mimic42.config import Settings

        settings = Settings()
        api_key = settings.openrouter_api_key
        if not api_key:
            return "[Ошибка: OPENROUTER_API_KEY не задан в окружении.]"

        audio_b64 = base64.b64encode(file_bytes).decode("utf-8")
        payload = {
            "model": "openai/whisper-large-v3",
            "input_audio": {
                "data": audio_b64,
                "format": filename.split(".")[-1],
            },
        }

        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(
                "https://openrouter.ai/api/v1/audio/transcriptions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json=payload,
            )
            if response.status_code == 200:
                res_json = response.json()
                return str(res_json.get("text", ""))
            else:
                return f"[Ошибка OpenRouter API (код {response.status_code}): {response.text}]"

    async def set_wakeup_timer(
        self, peer: str, delay_seconds: int, description: str
    ) -> dict[str, Any]:
        """Установить таймер пробуждения агента для выполнения отложенной задачи.

        Args:
            peer: Идентификатор чата, в котором сработает таймер (например, юзернейм или ID).
            delay_seconds: Задержка в секундах перед срабатыванием таймера.
            description: Описание задачи для напоминания при пробуждении.
        """
        if not self._agent_id or not self._session_factory:
            return {
                "success": False,
                "error": "Database session factory or agent_id not configured",
            }

        try:
            from datetime import UTC, datetime, timedelta

            from mimic42.integrations.database_models import AgentTimerModel

            trigger_at = datetime.now(UTC) + timedelta(seconds=delay_seconds)

            async with self._session_factory() as db_session:
                timer = AgentTimerModel(
                    agent_id=self._agent_id,
                    peer=peer,
                    trigger_at=trigger_at,
                    description=description,
                    status="pending",
                )
                db_session.add(timer)
                await db_session.commit()

            return {"success": True, "trigger_at": trigger_at.isoformat()}
        except Exception as e:
            return {"success": False, "error": str(e)}

    # Category 7: Chat Folder Management

    def _serialize_peer(self, peer: Any) -> dict[str, Any]:
        """Helper to serialize InputPeer/Peer objects into a simple dict."""
        if isinstance(peer, (types.InputPeerUser, types.PeerUser)):
            return {"type": "user", "id": getattr(peer, "user_id", None)}
        elif isinstance(peer, (types.InputPeerChannel, types.PeerChannel)):
            return {"type": "channel", "id": getattr(peer, "channel_id", None)}
        elif isinstance(peer, (types.InputPeerChat, types.PeerChat)):
            return {"type": "chat", "id": getattr(peer, "chat_id", None)}
        return {
            "type": "unknown",
            "id": getattr(peer, "id", None) or getattr(peer, "peer_id", None),
        }

    async def get_chat_folders(self) -> list[dict[str, Any]]:
        """Get user's chat folders (dialog filters)."""
        try:
            res = await self._client(functions.messages.GetDialogFiltersRequest())
            folders = []
            filters = res.filters if hasattr(res, "filters") else res
            for f in filters:
                if isinstance(f, types.DialogFilter):
                    folders.append(
                        {
                            "id": f.id,
                            "title": f.title.text if hasattr(f.title, "text") else str(f.title),
                            "emoticon": f.emoticon,
                            "color": f.color,
                            "pinned_peers": [
                                self._serialize_peer(p) for p in (f.pinned_peers or [])
                            ],
                            "include_peers": [
                                self._serialize_peer(p) for p in (f.include_peers or [])
                            ],
                            "exclude_peers": [
                                self._serialize_peer(p) for p in (f.exclude_peers or [])
                            ],
                            "contacts": bool(f.contacts),
                            "non_contacts": bool(f.non_contacts),
                            "groups": bool(f.groups),
                            "broadcasts": bool(f.broadcasts),
                            "bots": bool(f.bots),
                            "exclude_muted": bool(f.exclude_muted),
                            "exclude_read": bool(f.exclude_read),
                            "exclude_archived": bool(f.exclude_archived),
                        }
                    )
                elif isinstance(f, types.DialogFilterDefault):
                    folders.append({"id": 0, "title": "All Chats", "type": "default"})
            return folders
        except Exception as e:
            return [{"error": str(e)}]

    async def create_or_update_chat_folder(
        self,
        folder_id: int,
        title: str,
        emoticon: str | None = None,
        color: int | None = None,
        pinned_peers: list[str] | None = None,
        include_peers: list[str] | None = None,
        exclude_peers: list[str] | None = None,
        contacts: bool | None = None,
        non_contacts: bool | None = None,
        groups: bool | None = None,
        broadcasts: bool | None = None,
        bots: bool | None = None,
        exclude_muted: bool | None = None,
        exclude_read: bool | None = None,
        exclude_archived: bool | None = None,
    ) -> dict[str, Any]:
        """Create or update a custom chat folder (dialog filter)."""
        try:

            async def resolve_peers(peer_list):
                if not peer_list:
                    return []
                resolved = []
                for p in peer_list:
                    try:
                        entity = await self._resolve_peer(p)
                        resolved.append(entity)
                    except Exception:
                        pass
                return resolved

            pinned = await resolve_peers(pinned_peers)
            included = await resolve_peers(include_peers)
            excluded = await resolve_peers(exclude_peers)

            folder = types.DialogFilter(
                id=folder_id,
                title=title,
                pinned_peers=pinned,
                include_peers=included,
                exclude_peers=excluded,
                contacts=contacts,
                non_contacts=non_contacts,
                groups=groups,
                broadcasts=broadcasts,
                bots=bots,
                exclude_muted=exclude_muted,
                exclude_read=exclude_read,
                exclude_archived=exclude_archived,
                emoticon=emoticon,
                color=color,
            )

            await self._client(
                functions.messages.UpdateDialogFilterRequest(id=folder_id, filter=folder)
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def delete_chat_folder(self, folder_id: int) -> dict[str, Any]:
        """Delete a chat folder by its ID."""
        try:
            await self._client(
                functions.messages.UpdateDialogFilterRequest(id=folder_id, filter=None)
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    # Category 8: Privacy and Account Settings

    async def get_privacy_settings(self, key: PrivacyKey) -> dict[str, Any]:
        """Get privacy settings for a specific key."""
        try:
            key_cls = _PRIVACY_KEY_MAP.get(key)
            if not key_cls:
                return {"error": f"Unknown privacy key: {key}"}
            result = await self._client(functions.account.GetPrivacyRequest(key=key_cls()))
            rules = []
            for rule in result.rules:
                if isinstance(rule, types.PrivacyValueAllowAll):
                    rules.append("allow_all")
                elif isinstance(rule, types.PrivacyValueAllowContacts):
                    rules.append("allow_contacts")
                elif isinstance(rule, types.PrivacyValueAllowCloseFriends):
                    rules.append("allow_close_friends")
                elif isinstance(rule, types.PrivacyValueAllowPremium):
                    rules.append("allow_premium")
                elif isinstance(rule, types.PrivacyValueAllowBots):
                    rules.append("allow_bots")
                elif isinstance(rule, types.PrivacyValueAllowUsers):
                    rules.append({"type": "allow_users", "count": len(rule.users)})
                elif isinstance(rule, types.PrivacyValueAllowChatParticipants):
                    rules.append({"type": "allow_chats", "count": len(rule.chats)})
                elif isinstance(rule, types.PrivacyValueDisallowAll):
                    rules.append("disallow_all")
                elif isinstance(rule, types.PrivacyValueDisallowContacts):
                    rules.append("disallow_contacts")
                elif isinstance(rule, types.PrivacyValueDisallowBots):
                    rules.append("disallow_bots")
                elif isinstance(rule, types.PrivacyValueDisallowUsers):
                    rules.append({"type": "disallow_users", "count": len(rule.users)})
                elif isinstance(rule, types.PrivacyValueDisallowChatParticipants):
                    rules.append({"type": "disallow_chats", "count": len(rule.chats)})
                else:
                    rules.append({"type": type(rule).__name__})
            return {"key": key, "rules": rules}
        except Exception as e:
            return {"error": str(e)}

    async def set_privacy_settings(
        self,
        key: PrivacyKey,
        rule: PrivacyRule,
        allowed_users: list[str] | None = None,
        disallowed_users: list[str] | None = None,
    ) -> dict[str, Any]:
        """Set privacy settings for a specific key."""
        try:
            key_cls = _PRIVACY_KEY_MAP.get(key)
            if not key_cls:
                return {"success": False, "error": f"Unknown privacy key: {key}"}
            rule_cls = _PRIVACY_RULE_MAP.get(rule)
            if not rule_cls:
                return {"success": False, "error": f"Unknown privacy rule: {rule}"}
            tl_rules: list[types.TypeInputPrivacyRule] = [rule_cls()]
            if allowed_users:
                entities = []
                for u in allowed_users:
                    entities.append(await self._resolve_peer(u))
                tl_rules.append(types.InputPrivacyValueAllowUsers(users=entities))
            if disallowed_users:
                entities = []
                for u in disallowed_users:
                    entities.append(await self._resolve_peer(u))
                tl_rules.append(types.InputPrivacyValueDisallowUsers(users=entities))
            await self._client(
                functions.account.SetPrivacyRequest(
                    key=key_cls(),
                    rules=tl_rules,
                )
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def get_global_settings(self) -> dict[str, Any]:
        """Get global privacy settings."""
        try:
            res = await self._client(functions.account.GetGlobalPrivacySettingsRequest())
            return {
                "archive_and_mute_new_noncontact_peers": bool(
                    res.archive_and_mute_new_noncontact_peers
                ),
                "keep_archived_unmuted": bool(res.keep_archived_unmuted),
                "keep_archived_folders": bool(res.keep_archived_folders),
                "hide_read_marks": bool(res.hide_read_marks),
                "new_noncontact_peers_require_premium": bool(
                    res.new_noncontact_peers_require_premium
                ),
                "display_gifts_button": bool(res.display_gifts_button),
                "noncontact_peers_paid_stars": res.noncontact_peers_paid_stars,
            }
        except Exception as e:
            return {"error": str(e)}

    async def set_global_settings(
        self,
        archive_and_mute_new_noncontact_peers: bool | None = None,
        keep_archived_unmuted: bool | None = None,
        keep_archived_folders: bool | None = None,
        hide_read_marks: bool | None = None,
        new_noncontact_peers_require_premium: bool | None = None,
        display_gifts_button: bool | None = None,
    ) -> dict[str, Any]:
        """Set global privacy settings."""
        try:
            settings = types.GlobalPrivacySettings(
                archive_and_mute_new_noncontact_peers=archive_and_mute_new_noncontact_peers,
                keep_archived_unmuted=keep_archived_unmuted,
                keep_archived_folders=keep_archived_folders,
                hide_read_marks=hide_read_marks,
                new_noncontact_peers_require_premium=new_noncontact_peers_require_premium,
                display_gifts_button=display_gifts_button,
            )
            await self._client(functions.account.SetGlobalPrivacySettingsRequest(settings=settings))
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}

    async def get_content_settings(self) -> dict[str, Any]:
        """Get content settings (sensitive content filter)."""
        try:
            res = await self._client(functions.account.GetContentSettingsRequest())
            return {
                "sensitive_enabled": bool(res.sensitive_enabled),
                "sensitive_can_change": bool(res.sensitive_can_change),
            }
        except Exception as e:
            return {"error": str(e)}

    async def set_content_settings(self, sensitive_enabled: bool) -> dict[str, Any]:
        """Enable or disable sensitive content filter."""
        try:
            await self._client(
                functions.account.SetContentSettingsRequest(sensitive_enabled=sensitive_enabled)
            )
            return {"success": True}
        except Exception as e:
            return {"success": False, "error": str(e)}


def build_telegram_langchain_tools(
    client: TelethonRequestClient,
    agent_id: UUID | None = None,
    session_factory: async_sessionmaker[AsyncSession] | None = None,
) -> list[BaseTool]:
    """Expose all 89 tools as LangChain StructuredTools."""
    toolbox = TelegramToolbox(client, agent_id=agent_id, session_factory=session_factory)

    return [
        StructuredTool.from_function(
            coroutine=toolbox.send_text_message,
            name="send_text_message",
            description="Send a text message (markdown supported).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.edit_text_message,
            name="edit_text_message",
            description="Edit a message previously sent by the bot.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.delete_messages,
            name="delete_messages",
            description="Delete messages in a chat.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.forward_messages,
            name="forward_messages",
            description="Forward messages from one chat to another.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.pin_message,
            name="pin_message",
            description="Pin a message in a chat.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.unpin_message,
            name="unpin_message",
            description="Unpin a message in a chat.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.unpin_all_messages,
            name="unpin_all_messages",
            description="Unpin all pinned messages in a chat.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.send_chat_action,
            name="send_chat_action",
            description=(
                "Send a chat action indicator (typing, record_audio, "
                "upload_document, record_video)."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.send_reaction,
            name="send_reaction",
            description="Set or remove a reaction emoji on a message.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_message_reactions,
            name="get_message_reactions",
            description="Get the detailed reactions list of a message.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.mark_chat_as_read,
            name="mark_chat_as_read",
            description="Mark messages in a chat as read.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_messages,
            name="get_messages",
            description="Get message history (annotated with Media IDs: photo, sticker, doc).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_dialogs,
            name="get_dialogs",
            description="Get recent dialogs (chats, channels, users).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.search_messages,
            name="search_messages",
            description="Search messages inside a specific chat by keyword query.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.delete_dialog,
            name="delete_dialog",
            description="Delete a dialog (leaves a chat or channel).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.archive_dialogs,
            name="archive_dialogs",
            description="Archive dialogs.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.unarchive_dialogs,
            name="unarchive_dialogs",
            description="Unarchive dialogs.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_common_chats,
            name="get_common_chats",
            description="Get the list of common groups/channels with a user.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.send_file,
            name="send_file",
            description="Send a file or photo by URL or Media ID.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.view_image,
            name="view_image",
            description=(
                "View an image or sticker by its Media ID. Returns "
                "base64 image data in LangChain multimodal format."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.set_profile_photo,
            name="set_profile_photo",
            description="Set the bot's profile photo by URL or Media ID in memory.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.send_voice_note,
            name="send_voice_note",
            description="Send a voice note audio file by URL or Media ID.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.send_video_note,
            name="send_video_note",
            description="Send a video note (round video message) by URL or Media ID.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_sticker_sets,
            name="get_sticker_sets",
            description="Get user's installed sticker sets.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_stickers_in_set,
            name="get_stickers_in_set",
            description="Get list of stickers inside a specific set by set short name.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.search_sticker_sets,
            name="search_sticker_sets",
            description="Search sticker sets globally by name query.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.install_sticker_set,
            name="install_sticker_set",
            description="Install/add a sticker set to user's list.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.uninstall_sticker_set,
            name="uninstall_sticker_set",
            description="Uninstall/remove a sticker set from user's list.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.send_sticker,
            name="send_sticker",
            description="Send a sticker by its Media ID.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_profile,
            name="get_profile",
            description=(
                "Get high-level details of a user: bio/about, name, "
                "online status, common chats count."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.update_profile_info,
            name="update_profile_info",
            description="Update bot's name and bio (about).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.update_username,
            name="update_username",
            description="Update bot's public @username.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.add_contact,
            name="add_contact",
            description="Add a contact by phone number.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.delete_contact,
            name="delete_contact",
            description="Delete user from contact list.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_contacts,
            name="get_contacts",
            description="Get contacts list.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_chat_info,
            name="get_chat_info",
            description=(
                "Get high-level details of a group/channel (about, title, participants count)."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.check_admin_permissions,
            name="check_admin_permissions",
            description="Check administrative rights of the bot itself in a group or channel.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.create_group,
            name="create_group",
            description="Create a simple group chat with list of user IDs/usernames.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.create_channel,
            name="create_channel",
            description="Create a channel or supergroup.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.invite_to_channel,
            name="invite_to_channel",
            description="Invite/add users to channel/supergroup.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.kick_chat_member,
            name="kick_chat_member",
            description="Kick/remove participant from chat.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.ban_chat_member,
            name="ban_chat_member",
            description="Ban participant in group/channel.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.restrict_chat_member,
            name="restrict_chat_member",
            description="Restrict permissions of participant in chat.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.promote_chat_member,
            name="promote_chat_member",
            description="Promote participant to administrator with custom title (member tag).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_chat_members,
            name="get_chat_members",
            description="Get members of a group/channel (with custom titles).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_chat_admin_log,
            name="get_chat_admin_log",
            description="Get administrative audit log.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.edit_chat_title,
            name="edit_chat_title",
            description="Change the title of a channel, group or supergroup.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.edit_chat_about,
            name="edit_chat_about",
            description="Change the about/description of a channel, group or supergroup.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.edit_chat_photo,
            name="edit_chat_photo",
            description=(
                "Change the photo of a channel, group or supergroup. "
                "photo can be a file path or URL."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.update_chat_public_link,
            name="update_chat_public_link",
            description=(
                "Set or change the public username/link of a channel or supergroup. "
                "Pass empty string to remove."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.set_chat_default_banned_rights,
            name="set_chat_default_banned_rights",
            description=(
                "Set default restricted rights for all members of a group/channel. "
                "rights is a JSON string mapping flag names to true/false."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.toggle_chat_signatures,
            name="toggle_chat_signatures",
            description="Toggle message signatures in a channel (shows admin name on posts).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.delete_channel,
            name="delete_channel",
            description="Delete a channel or supergroup entirely.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.toggle_join_requests,
            name="toggle_join_requests",
            description="Enable or disable join requests (approval required to join).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.toggle_join_to_send,
            name="toggle_join_to_send",
            description=(
                "Enable or disable the requirement to join the channel "
                "before sending messages."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.toggle_slow_mode,
            name="toggle_slow_mode",
            description=(
                "Enable or disable slow mode in a group. "
                "seconds can be 0 (off), 10, 30, 60, 300, 900, 3600."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.set_discussion_group,
            name="set_discussion_group",
            description="Link a discussion group (supergroup) to a broadcast channel.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.toggle_forum,
            name="toggle_forum",
            description="Enable or disable forum mode in a supergroup (creates topics).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.toggle_pre_history_hidden,
            name="toggle_pre_history_hidden",
            description="Hide or show previous chat history for new members.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.toggle_participants_hidden,
            name="toggle_participants_hidden",
            description="Hide or show the list of participants in a channel/group.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.edit_chat_location,
            name="edit_chat_location",
            description="Set a geolocation for a group/channel (appears in info).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.toggle_anti_spam,
            name="toggle_anti_spam",
            description="Enable or disable native Telegram anti-spam protection in a group.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.set_chat_admin_rights,
            name="set_chat_admin_rights",
            description=(
                "Set full admin rights for a user in a channel/group. "
                "rights is a JSON string mapping flag names to true/false."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.set_chat_banned_rights,
            name="set_chat_banned_rights",
            description=(
                "Set full banned/restricted rights for a user in a channel/group. "
                "rights is a JSON string mapping flag names to true/false."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.join_channel,
            name="join_channel",
            description="Join a public channel/group.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.send_poll,
            name="send_poll",
            description="Send a poll or quiz.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.transcribe_voice_note,
            name="transcribe_voice_note",
            description="Transcribe a voice note or round video message to text.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.read_document_file,
            name="read_document_file",
            description="Read and extract text/contents of a document or file.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.set_wakeup_timer,
            name="set_wakeup_timer",
            description=(
                "Set a wakeup timer to trigger a delayed task in a chat after a delay in seconds."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.mute_chat,
            name="mute_chat",
            description="Mute a chat. If duration_hours is not specified, it is muted indefinitely (10 years).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.unmute_chat,
            name="unmute_chat",
            description="Unmute a chat to enable normal message notifications.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.send_location,
            name="send_location",
            description="Send a map location pin with specific latitude and longitude coordinates.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.send_venue,
            name="send_venue",
            description="Send a beautiful venue location card with a map pin, title, and address.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.search_location,
            name="search_location",
            description="Search for a location or address and return its coordinates and formatted address using ArcGIS.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_chat_folders,
            name="get_chat_folders",
            description="Get user's chat folders (dialog filters).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.create_or_update_chat_folder,
            name="create_or_update_chat_folder",
            description="Create or update a custom chat folder (dialog filter).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.delete_chat_folder,
            name="delete_chat_folder",
            description="Delete a custom chat folder by its ID.",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_message_buttons,
            name="get_message_buttons",
            description=(
                "Get inline or reply keyboard buttons from a message. "
                "CRITICAL: The 'peer' argument MUST be the chat/bot where the message is located, not the current chat! "
                "Returns list of buttons with text, type, data, URL, etc."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.click_inline_button,
            name="click_inline_button",
            description=(
                "Click an inline callback button on a message. "
                "Use button_data (bytes as string) or button_index to identify the button."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.click_reply_keyboard_button,
            name="click_reply_keyboard_button",
            description=(
                "Press a reply keyboard (markup) button by sending its text as a message."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.query_inline_bot,
            name="query_inline_bot",
            description=(
                "Query an inline bot (e.g., @vote query) and return results with query_id."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.send_inline_bot_result,
            name="send_inline_bot_result",
            description=("Send an inline bot result to a chat using query_id and result_id."),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.start_bot,
            name="start_bot",
            description=("Start a bot with optional deep-link parameter (e.g., /start ref123)."),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_privacy_settings,
            name="get_privacy_settings",
            description=(
                "Get privacy settings for a specific key. "
                "Keys: status_timestamp, profile_photo, phone_number, added_by_phone, "
                "chat_invite, phone_call, phone_p2p, forwards, voice_messages, "
                "about, birthday, saved_music, no_paid_messages, star_gifts_auto_save."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.set_privacy_settings,
            name="set_privacy_settings",
            description=(
                "Set privacy settings for a specific key. "
                "Rules: allow_all, allow_contacts, allow_close_friends, allow_premium, "
                "allow_bots, disallow_all, disallow_contacts, disallow_bots. "
                "Optionally add allowed_users or disallowed_users for exceptions."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_global_settings,
            name="get_global_settings",
            description="Get global privacy settings (archive, read marks, etc.).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.set_global_settings,
            name="set_global_settings",
            description=(
                "Set global privacy settings: archive_and_mute_new_noncontact_peers, "
                "keep_archived_unmuted, keep_archived_folders, hide_read_marks, "
                "new_noncontact_peers_require_premium, display_gifts_button."
            ),
        ),
        StructuredTool.from_function(
            coroutine=toolbox.get_content_settings,
            name="get_content_settings",
            description="Get content settings (sensitive content filter status).",
        ),
        StructuredTool.from_function(
            coroutine=toolbox.set_content_settings,
            name="set_content_settings",
            description="Enable or disable sensitive content filter.",
        ),
    ]
