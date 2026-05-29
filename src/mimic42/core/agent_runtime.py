from __future__ import annotations

import asyncio
import logging
import random
from collections.abc import Awaitable, Callable, Mapping
from enum import StrEnum
from typing import Any, Protocol, cast
from uuid import UUID

from pydantic import BaseModel, Field
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker

from mimic42.core.memory import MemoryServiceLike, RuntimeMemoryService

logger = logging.getLogger("mimic42.agent_runtime")


class TelegramAuthorizationRequired(RuntimeError):
    """Raised when a Telethon user session is connected but not authorized."""


class AgentRuntimeState(StrEnum):
    STOPPED = "stopped"
    STARTING = "starting"
    RUNNING = "running"
    STOPPING = "stopping"
    ERROR = "error"


class AgentRuntimeConfig(BaseModel):
    agent_id: UUID
    owner_id: UUID
    telegram_session_name: str = Field(min_length=1)
    telegram_api_id: int = Field(gt=0)
    telegram_api_hash: str = Field(min_length=1)
    telegram_session_string: str | None = Field(default=None, min_length=1)
    llm_model: str = Field(default="mistralai/mistral-small-2603", min_length=1)
    system_prompt: str = Field(min_length=1)
    soul_prompt: str = Field(default="", max_length=20_000)
    name: str = Field(default="AI", min_length=1, max_length=120)

    @property
    def combined_prompt(self) -> str:
        prompt = self.system_prompt
        prompt = prompt.replace("{{name}}", self.name)
        prompt = prompt.replace("{{soul}}", self.soul_prompt)
        return prompt


class AgentStatus(BaseModel):
    agent_id: UUID
    owner_id: UUID
    state: AgentRuntimeState


class AgentTrigger(BaseModel):
    peer: str = Field(min_length=1)
    text: str = Field(min_length=1)
    message_id: int | None = Field(default=None, gt=0)


class AgentTriggerResult(BaseModel):
    agent_id: UUID
    peer: str
    input_text: str
    response_text: str
    telegram_message_id: str | None = None


class TelegramClientLike(Protocol):
    async def connect(self) -> None: ...

    async def disconnect(self) -> None: ...

    async def is_user_authorized(self) -> bool: ...

    async def send_message(self, entity: str, message: str, **kwargs: Any) -> object: ...

    def add_event_handler(
        self,
        callback: Callable[[object], Awaitable[None]],
        event: object | None = None,
    ) -> None: ...


class LangChainAgentLike(Protocol):
    async def ainvoke(self, input_data: dict[str, object]) -> object: ...


class MimicAgentRuntime:
    """Async runtime that owns one Telegram user session and one LangChain agent."""

    def __init__(
        self,
        *,
        config: AgentRuntimeConfig,
        telegram_client: TelegramClientLike,
        langchain_agent: LangChainAgentLike,
        memory_service: MemoryServiceLike | None = None,
        session_factory: async_sessionmaker[AsyncSession] | None = None,
    ) -> None:
        self.config = config
        self._telegram_client = telegram_client
        self._langchain_agent = langchain_agent
        self._memory_service = memory_service or RuntimeMemoryService()
        self._session_factory = session_factory
        self._state = AgentRuntimeState.STOPPED
        self._lifecycle_lock = asyncio.Lock()
        self._trigger_lock = asyncio.Lock()
        self._message_handler_registered = False
        self._member_tag_cache: dict[tuple[int, int], tuple[str | None, float]] = {}
        self._chat_mute_cache: dict[str, tuple[bool, float]] = {}
        self._scheduler_task: asyncio.Task[None] | None = None
        self._http_client: Any | None = None

    @property
    def state(self) -> AgentRuntimeState:
        return self._state

    @property
    def status(self) -> AgentStatus:
        return AgentStatus(
            agent_id=self.config.agent_id,
            owner_id=self.config.owner_id,
            state=self.state,
        )

    async def start(self) -> None:
        async with self._lifecycle_lock:
            if self._state is AgentRuntimeState.RUNNING:
                logger.debug(f"Agent {self.config.agent_id} already running")
                return

            logger.info(f"Starting agent {self.config.agent_id}")
            self._state = AgentRuntimeState.STARTING
            try:
                logger.debug("Connecting to Telegram...")
                await self._telegram_client.connect()
                logger.debug("Connected. Checking authorization...")
                if not await self._telegram_client.is_user_authorized():
                    logger.error("Telegram session not authorized")
                    raise TelegramAuthorizationRequired(
                        "Telegram user session is not authorized. Complete onboarding first."
                    )
                logger.debug("Authorized. Registering message handler...")
                self._register_message_handler()
                logger.info("Message handler registered")
            except Exception as e:
                logger.error(f"Failed to start agent {self.config.agent_id}: {e}", exc_info=True)
                self._state = AgentRuntimeState.ERROR
                raise

            self._state = AgentRuntimeState.RUNNING
            if self._session_factory is not None:
                self._scheduler_task = asyncio.create_task(self._run_scheduler_loop())
            import httpx
            self._http_client = httpx.AsyncClient(timeout=30.0)
            logger.info(f"Agent {self.config.agent_id} started successfully")

    async def stop(self) -> None:
        async with self._lifecycle_lock:
            if self._state is AgentRuntimeState.STOPPED:
                return

            self._state = AgentRuntimeState.STOPPING
            if self._scheduler_task is not None:
                self._scheduler_task.cancel()
                try:
                    await self._scheduler_task
                except asyncio.CancelledError:
                    pass
                self._scheduler_task = None

            if self._http_client is not None:
                await self._http_client.aclose()
                self._http_client = None

            await self._telegram_client.disconnect()
            self._state = AgentRuntimeState.STOPPED

    async def _humanized_send(
        self,
        peer: Any,
        text: str,
        reply_to: int | None = None,
    ) -> Any:
        """Send a message after a human-like typing delay with intermittent typing indicators."""
        from telethon import functions, types

        # Resolve peer entity for typing actions
        entity = peer
        get_input_entity = getattr(self._telegram_client, "get_input_entity", None)
        if callable(get_input_entity):
            try:
                entity = await get_input_entity(peer)
            except Exception:
                pass

        # Calculate typing delay: ~14 chars per second + random variance
        base_delay = len(text) * 0.07
        delay = max(0.8, min(15.0, base_delay + random.uniform(0.2, 1.2)))
        logger.debug("Humanized send: calculated delay %.2fs for %d chars", delay, len(text))

        end_time = asyncio.get_event_loop().time() + delay

        # Start typing
        try:
            await self._telegram_client(
                functions.messages.SetTypingRequest(
                    peer=entity,
                    action=types.SendMessageTypingAction(),
                )
            )
        except Exception:
            logger.debug("Failed to start typing action", exc_info=True)

        while True:
            remaining = end_time - asyncio.get_event_loop().time()
            if remaining <= 0:
                break
            sleep_for = min(0.5, remaining)
            await asyncio.sleep(sleep_for)

            # Small chance (8%) to briefly interrupt typing for realism
            if random.random() < 0.08:
                try:
                    await self._telegram_client(
                        functions.messages.SetTypingRequest(
                            peer=entity,
                            action=types.SendMessageCancelAction(),
                        )
                    )
                    await asyncio.sleep(random.uniform(0.05, 0.25))
                    await self._telegram_client(
                        functions.messages.SetTypingRequest(
                            peer=entity,
                            action=types.SendMessageTypingAction(),
                        )
                    )
                except Exception:
                    pass
            else:
                try:
                    await self._telegram_client(
                        functions.messages.SetTypingRequest(
                            peer=entity,
                            action=types.SendMessageTypingAction(),
                        )
                    )
                except Exception:
                    pass

        # Send the actual message
        try:
            if reply_to is not None:
                sent_message = await self._telegram_client.send_message(
                    peer,
                    text,
                    reply_to=reply_to,
                )
            else:
                sent_message = await self._telegram_client.send_message(peer, text)
        finally:
            # Cancel typing
            try:
                await self._telegram_client(
                    functions.messages.SetTypingRequest(
                        peer=entity,
                        action=types.SendMessageCancelAction(),
                    )
                )
            except Exception:
                pass

        return sent_message

    async def trigger_message(self, trigger: AgentTrigger) -> AgentTriggerResult:
        logger.debug(
            "trigger_message called: peer=%s, text=%s",
            trigger.peer,
            trigger.text[:50],
        )
        if self._state is not AgentRuntimeState.RUNNING:
            logger.info(f"Agent not running (state={self._state}), starting...")
            await self.start()

        async with self._trigger_lock:
            logger.debug(f"Processing message from {trigger.peer}: {trigger.text[:100]}")
            messages = await self._memory_service.build_messages(
                agent_id=self.config.agent_id,
                peer=trigger.peer,
                user_text=trigger.text,
            )
            logger.debug(f"Built {len(messages)} messages for context")
            try:
                response = await self._langchain_agent.ainvoke(
                    {
                        "messages": messages,
                    }
                )
                logger.debug(f"Agent response: {response}")
            except Exception as e:
                logger.error(f"Error invoking agent: {e}", exc_info=True)
                raise

            # Convert output BaseMessage objects to plain dicts for storage/comparison.
            output_messages = _messages_to_dicts(response)

            # Interpret structured response if provided by the agent. Expected format:
            # { "send_any_message": bool, "text": str, "reply_to": int|str }
            send_any, response_text, reply_to = _interpret_agent_response(response)
            display_text = response_text[:100] if response_text else "None"
            logger.debug(f"Interpreted response: send_any={send_any}, text={display_text}")

            # Don't send empty messages
            if send_any and not response_text:
                logger.debug("Agent generated empty response, not sending")
                send_any = False

            # Convert stringified numeric peer ID to integer for Telethon compatibility
            peer_id_value: str | int = trigger.peer
            if isinstance(peer_id_value, str):
                if peer_id_value.startswith("-") and peer_id_value[1:].isdigit():
                    peer_id_value = int(peer_id_value)
                elif peer_id_value.isdigit():
                    peer_id_value = int(peer_id_value)
            # Use the correctly typed value for sending
            peer_id_for_send = peer_id_value
            # Mark incoming message as read immediately after deciding to reply,
            # before the typing delay, so the order is: read -> typing -> send.
            if send_any and trigger.message_id is not None:
                try:
                    from telethon.tl import functions
                    read_entity = peer_id_for_send
                    get_input_entity = getattr(self._telegram_client, "get_input_entity", None)
                    if callable(get_input_entity):
                        try:
                            read_entity = await get_input_entity(peer_id_for_send)
                        except Exception:
                            pass
                    await self._telegram_client(
                        functions.messages.ReadHistoryRequest(
                            peer=read_entity,
                            max_id=trigger.message_id,
                        )
                    )
                    logger.debug(
                        "Marked message %s as read in peer %s",
                        trigger.message_id,
                        trigger.peer,
                    )
                except Exception:
                    logger.warning("Failed to mark message as read", exc_info=True)

            sent_message = None
            if send_any:
                logger.info(f"Sending response to {peer_id_for_send}: {response_text[:100]}")
                try:
                    sent_message = await self._humanized_send(
                        peer_id_for_send,
                        response_text,
                        reply_to=reply_to,
                    )
                    logger.info(f"Message sent successfully to {peer_id_for_send}")
                except Exception:
                    logger.exception("Failed to send Telegram message to %s", peer_id_for_send)
            else:
                logger.debug("Agent decided not to send message (send_any=False)")

            await self._memory_service.save_messages(
                agent_id=self.config.agent_id,
                peer=trigger.peer,
                input_messages=messages,
                output_messages=output_messages,
            )

        return AgentTriggerResult(
            agent_id=self.config.agent_id,
            peer=trigger.peer,
            input_text=trigger.text,
            response_text=response_text,
            telegram_message_id=(
                _extract_message_id(sent_message) if sent_message is not None else None
            ),
        )

    def _register_message_handler(self) -> None:
        if self._message_handler_registered:
            return

        try:
            from telethon import events
        except ImportError:
            event_builder = None
        else:
            event_builder = events.NewMessage(incoming=True)

        self._telegram_client.add_event_handler(self._handle_incoming_message, event_builder)
        self._message_handler_registered = True

    async def _handle_incoming_message(self, event: object) -> None:
        logger.debug("Incoming message event received")
        logger.debug(
            "Incoming message event: chat_id=%s, text=%s",
            getattr(event, "chat_id", None),
            getattr(event, "raw_text", "")[:50],
        )
        # Check if chat is muted
        try:
            peer = await _extract_incoming_peer(event)
            import time
            now_ts = time.time()
            
            is_muted = False
            if peer in self._chat_mute_cache:
                cached_muted, expiry = self._chat_mute_cache[peer]
                if now_ts < expiry:
                    is_muted = cached_muted
                    if is_muted:
                        return
                else:
                    self._chat_mute_cache.pop(peer, None)
                    
            if peer not in self._chat_mute_cache:
                input_chat = None
                get_input_chat = getattr(event, "get_input_chat", None)
                if callable(get_input_chat):
                    try:
                        input_chat = await event.get_input_chat()
                    except Exception:
                        logger.warning("Failed to get input chat for mute check", exc_info=True)
                
                if input_chat is None:
                    input_chat = getattr(event, "input_chat", None)
                
                if input_chat is None:
                    get_input_entity = getattr(self._telegram_client, "get_input_entity", None)
                    if callable(get_input_entity):
                        input_chat = await get_input_entity(peer)
                
                if input_chat is not None:
                    from telethon import functions, types
                    notify_peer = types.InputNotifyPeer(peer=input_chat)
                    res = await self._telegram_client(
                        functions.account.GetNotifySettingsRequest(peer=notify_peer)
                    )
                    
                    is_muted = False
                    if getattr(res, "silent", False):
                        is_muted = True
                    if getattr(res, "mute_until", None):
                        from datetime import datetime
                        if res.mute_until.tzinfo:
                            now = datetime.now(res.mute_until.tzinfo)
                        else:
                            now = datetime.now()
                        if res.mute_until > now:
                            is_muted = True
                    
                    self._chat_mute_cache[peer] = (is_muted, now_ts + 60.0)
                    
                    if is_muted:
                        logger.info("Chat %s is muted, skipping", peer)
                        return
        except Exception:
            logger.exception("Failed to check mute status for peer %s", peer)

        # Protect the rest of the message handling pipeline from crashes
        try:
            raw_text = getattr(event, "raw_text", None) or getattr(event, "text", None)
            if not isinstance(raw_text, str):
                raw_text = ""

            # Process attachments in memory and transcribers
            text = await _process_media_and_text(event, raw_text, http_client=self._http_client)
            if not text:
                logger.info(
                    "Empty text after _process_media_and_text for chat %s, skipping",
                    getattr(event, "chat_id", None),
                )
                return

            # Format sender name and metadata
            is_private = getattr(event, "is_private", False)
            is_group = getattr(event, "is_group", False)

            from datetime import datetime
            msg_date = getattr(event, "date", None)
            if not msg_date:
                message = getattr(event, "message", None)
                msg_date = getattr(message, "date", None)
            if not msg_date:
                msg_date = datetime.now()
            time_str = msg_date.strftime("%Y-%m-%d %H:%M:%S")

            # Chat name
            if is_private:
                chat_type_str = "ЛС"
            else:
                chat = await event.get_chat()
                chat_title = getattr(chat, "title", "")
                if not chat_title:
                    chat_title = (
                        getattr(chat, "username", "")
                        or str(getattr(event, "chat_id", ""))
                    )
                if is_group:
                    chat_type_str = f'Группа "{chat_title}"'
                else:
                    chat_type_str = f'Канал "{chat_title}"'

            # Sender details
            get_sender = getattr(event, "get_sender", None)
            sender = None
            if callable(get_sender):
                try:
                    import inspect
                    res = get_sender()
                    if inspect.isawaitable(res):
                        sender = await res
                    else:
                        sender = res
                except Exception:
                    logger.warning("Failed to get sender for event", exc_info=True)
            if sender:
                first_name = getattr(sender, "first_name", None) or ""
                last_name = getattr(sender, "last_name", None) or ""
                name_parts = []
                if first_name:
                    name_parts.append(first_name)
                if last_name:
                    name_parts.append(last_name)
                name_str = " ".join(name_parts)
                if not name_str:
                    name_str = (
                        getattr(sender, "title", None)
                        or getattr(sender, "username", None)
                        or str(getattr(sender, "id", ""))
                    )
                if not name_str:
                    name_str = "Unknown"

                username = getattr(sender, "username", None)
                username_str = f"@{username}" if username else ""
                sender_id = getattr(sender, "id", None)
                id_str = f"ID: {sender_id}" if sender_id else ""

                details = ", ".join(filter(None, [username_str, id_str]))
                details_str = f" ({details})" if details else ""
                sender_str = f"{name_str}{details_str}"
            else:
                chat = await event.get_chat()
                chat_title = getattr(chat, "title", "Unknown")
                sender_str = chat_title

            # Check role/title
            title = None
            if getattr(event, "sender_id", None) and getattr(event, "chat_id", None):
                chat_id = event.chat_id
                sender_id = event.sender_id
                cache_key = (chat_id, sender_id)
                import time
                now_ts = time.time()
                if cache_key in self._member_tag_cache:
                    cached_title, expiry = self._member_tag_cache[cache_key]
                    if now_ts < expiry:
                        title = cached_title

                is_expired = (
                    cache_key not in self._member_tag_cache
                    or now_ts >= self._member_tag_cache[cache_key][1]
                )
                if is_expired:
                    try:
                        from telethon.tl import functions
                        is_supergroup = getattr(event, "is_channel", False)
                        if is_supergroup:
                            input_chat = getattr(event, "input_chat", None) or chat_id
                            input_sender = (
                                getattr(event, "input_sender", None) or sender_id
                            )
                            res = await event.client(
                                functions.channels.GetParticipantRequest(
                                    channel=input_chat,
                                    participant=input_sender,
                                )
                            )
                            title = (
                                res.participant.title
                                if hasattr(res.participant, "title")
                                else None
                            )
                    except Exception:
                        logger.warning("Failed to get participant title", exc_info=True)
                        title = None
                    self._member_tag_cache[cache_key] = (title, now_ts + 3600.0)

            # Fallback to channel post author signature
            post_author = getattr(getattr(event, "message", None), "post_author", None)
            if not title and post_author:
                title = post_author

            if title:
                sender_str += f" [Подпись/Роль: {title}]"

            # NewMessage.Event delegates __getattr__ to self.message, but
            # self.message is a raw types.Message (not the custom wrapper),
            # so it lacks the reply_to_msg_id property. Inspect reply_to directly.
            from telethon.tl import types

            reply_to_msg_id = None
            ev_message = getattr(event, "message", None)
            if ev_message:
                reply_to = getattr(ev_message, "reply_to", None)
                if isinstance(reply_to, types.MessageReplyHeader):
                    reply_to_msg_id = reply_to.reply_to_msg_id
                elif reply_to:
                    reply_to_msg_id = getattr(reply_to, "reply_to_msg_id", None)

            reply_str = ""
            if reply_to_msg_id:
                logger.debug(
                    "Reply detected: reply_to_msg_id=%s for chat_id=%s",
                    reply_to_msg_id,
                    getattr(event, "chat_id", None),
                )
                reply_preview = ""
                try:
                    reply_msg = await event.get_reply_message()
                    if reply_msg:
                        raw = getattr(reply_msg, "raw_text", "")
                        reply_preview = raw or getattr(reply_msg, "text", "")
                        logger.debug(
                            "get_reply_message succeeded, preview=%s",
                            reply_preview[:30] if reply_preview else "(empty)",
                        )
                    else:
                        logger.debug("get_reply_message returned None")
                except Exception:
                    logger.debug("get_reply_message failed", exc_info=True)

                if not reply_preview:
                    # Fallback: fetch the replied message directly via RPC
                    try:
                        get_messages = getattr(self._telegram_client, "get_messages", None)
                        if callable(get_messages):
                            peer = await _extract_incoming_peer(event)
                            msgs = await get_messages(
                                peer, ids=reply_to_msg_id
                            )
                            if msgs:
                                reply_msg = msgs[0] if isinstance(msgs, list) else msgs
                                raw = getattr(reply_msg, "raw_text", "")
                                reply_preview = raw or getattr(reply_msg, "text", "")
                                logger.debug(
                                    "Fallback get_messages succeeded, preview=%s",
                                    reply_preview[:30] if reply_preview else "(empty)",
                                )
                    except Exception:
                        logger.debug("Fallback get_messages failed", exc_info=True)

                if reply_preview:
                    preview = reply_preview[:20]
                    reply_str = (
                        f'Ответ на сообщение #{reply_to_msg_id}'
                        f' ("{preview}...")\n'
                    )
                else:
                    reply_str = f"Ответ на сообщение #{reply_to_msg_id}\n"

            # Format output message text
            text = (
                f"[Входящее сообщение]\n"
                f"Время: {time_str}\n"
                f"Чат: {chat_type_str}\n"
                f"Отправитель: {sender_str}\n"
                f"{reply_str}"
                f"Содержимое: {text}"
            )

            peer = await _extract_incoming_peer(event)
            await self.trigger_message(
                AgentTrigger(
                    peer=peer,
                    text=text,
                    message_id=_extract_incoming_message_id(event),
                )
            )
        except Exception:
            logger.exception("Unhandled exception in incoming message handler")

    async def _run_scheduler_loop(self) -> None:
        """Background loop to check and trigger pending agent timers."""
        while self._state is AgentRuntimeState.RUNNING:
            try:
                await self._check_and_trigger_timers()
            except asyncio.CancelledError:
                break
            except Exception:
                logger.exception("Error in scheduler loop")

            try:
                await asyncio.sleep(5.0)
            except asyncio.CancelledError:
                break

    async def _check_and_trigger_timers(self) -> None:
        if not self._session_factory:
            return

        from datetime import UTC, datetime

        from sqlalchemy import select, update

        from mimic42.integrations.database_models import AgentTimerModel

        now_utc = datetime.now(UTC)

        async with self._session_factory() as db_session:
            stmt = (
                select(AgentTimerModel)
                .where(
                    AgentTimerModel.agent_id == self.config.agent_id,
                    AgentTimerModel.status == "pending",
                    AgentTimerModel.trigger_at <= now_utc,
                )
            )
            result = await db_session.scalars(stmt)
            due_timers = list(result)

            if not due_timers:
                return

            for timer in due_timers:
                timer.status = "running"
            await db_session.commit()

            for timer in due_timers:
                try:
                    trigger_text = f"[Отложенное событие] Сработал таймер: {timer.description}"
                    await self.trigger_message(
                        AgentTrigger(
                            peer=timer.peer,
                            text=trigger_text,
                        )
                    )
                    timer.status = "succeeded"
                except Exception:
                    logger.exception("Error triggering timer %s", timer.id)
                    timer.status = "failed"

                # Update status
                async with self._session_factory() as update_session:
                    await update_session.execute(
                        update(AgentTimerModel)
                        .where(AgentTimerModel.id == timer.id)
                        .values(status=timer.status)
                    )
                    await update_session.commit()


async def _process_media_and_text(
    event: object,
    text: str,
    *,
    http_client: Any | None = None,
) -> str:
    message = getattr(event, "message", None)
    if not message or not getattr(message, "media", None):
        return text

    try:
        from mimic42.integrations.telegram_tools import format_media_object
        media_id = format_media_object(message)
        if not media_id:
            return text

        if media_id.startswith("photo:"):
            return f"[Фото id={media_id}]" + (f" {text}" if text else "")

        elif media_id.startswith("sticker:"):
            parts = media_id.split(":")
            emoji = parts[5] if len(parts) > 5 else ""
            pack_name = parts[6] if len(parts) > 6 else ""
            pack_str = f" пак={pack_name}" if pack_name else ""
            return (
                f"[Стикер {emoji} id={media_id}{pack_str}]"
                + (f" {text}" if text else "")
            )

        elif media_id.startswith(("voice:", "round:")):
            import os
            from io import BytesIO

            import httpx
            from mimic42.config import Settings

            settings = Settings()
            api_key = settings.openrouter_api_key
            if not api_key:
                err_msg = "[Голосовое сообщение (ошибка: OPENROUTER_API_KEY не установлен)]"
                return err_msg + (f" {text}" if text else "")

            buffer = BytesIO()
            await event.client.download_media(message, file=buffer)
            file_bytes = buffer.getvalue()
            if not file_bytes:
                err_msg = "[Голосовое сообщение (ошибка: файл пустой)]"
                return err_msg + (f" {text}" if text else "")

            filename = "voice.ogg" if media_id.startswith("voice:") else "video.mp4"

            try:
                client = http_client
                if client is None:
                    client = httpx.AsyncClient(timeout=30.0)
                import base64
                audio_b64 = base64.b64encode(file_bytes).decode("utf-8")
                payload = {
                    "model": "openai/whisper-large-v3",
                    "input_audio": {
                        "data": audio_b64,
                        "format": filename.split(".")[-1],
                    },
                }
                headers = {
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                }
                response = await client.post(
                    "https://openrouter.ai/api/v1/audio/transcriptions",
                    headers=headers,
                    json=payload,
                )
                response.raise_for_status()
                res_json = response.json()
                transcription = res_json.get("text", "")
                mtype = (
                    "Голосовое сообщение"
                    if media_id.startswith("voice:")
                    else "Видеосообщение"
                )
                trans_text = f"[{mtype} (расшифровка: \"{transcription}\")]"
                return trans_text + (f" {text}" if text else "")
            except Exception as e:
                mtype = (
                    "Голосовое сообщение"
                    if media_id.startswith("voice:")
                    else "Видеосообщение"
                )
                return f"[{mtype} (ошибка транскрипции: {e})]" + (f" {text}" if text else "")

        elif media_id.startswith("doc:"):
            parts = media_id.split(":")
            filename = parts[5] if len(parts) > 5 else "file"
            ext = filename.split(".")[-1].lower() if "." in filename else ""

            allowed_exts = (
                "docx",
                "xlsx",
                "txt",
                "md",
                "json",
                "csv",
                "xml",
                "py",
                "html",
                "css",
                "yaml",
                "yml",
            )
            if ext not in allowed_exts and ext != "":
                return (
                    f"[Файл name={filename} (этот тип документа нельзя открыть)]"
                    + (f" {text}" if text else "")
                )

            from io import BytesIO
            buffer = BytesIO()
            await event.client.download_media(message, file=buffer)
            file_bytes = buffer.getvalue()
            if not file_bytes:
                return f"[Файл name={filename} (пустой)]" + (f" {text}" if text else "")

            if ext == "docx":
                try:
                    import docx
                    doc = docx.Document(BytesIO(file_bytes))
                    paragraphs = [p.text for p in doc.paragraphs]
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text for cell in row.cells]
                            paragraphs.append(" | ".join(row_text))
                    doc_content = "\n".join(paragraphs)
                    doc_text = f"[Файл name={filename} (содержимое: \"{doc_content}\")]"
                    return doc_text + (f" {text}" if text else "")
                except Exception as e:
                    err_msg = f"[Файл name={filename} (ошибка чтения: {e})]"
                    return err_msg + (f" {text}" if text else "")

            elif ext == "xlsx":
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(BytesIO(file_bytes), read_only=True)
                    sheet_texts = []
                    for sheet in wb.worksheets:
                        sheet_texts.append(f"Лист: {sheet.title}")
                        for row in sheet.iter_rows(values_only=True):
                            row_str = " | ".join(str(val) if val is not None else "" for val in row)
                            if row_str.strip(" |"):
                                sheet_texts.append(row_str)
                    xlsx_content = "\n".join(sheet_texts)
                    xlsx_text = f"[Файл name={filename} (содержимое: \"{xlsx_content}\")]"
                    return xlsx_text + (f" {text}" if text else "")
                except Exception as e:
                    err_msg = f"[Файл name={filename} (ошибка чтения: {e})]"
                    return err_msg + (f" {text}" if text else "")

            else:
                try:
                    txt_content = file_bytes.decode("utf-8")
                except UnicodeDecodeError:
                    try:
                        txt_content = file_bytes.decode("latin-1")
                    except Exception as e:
                        txt_content = f"<Ошибка декодирования: {e}>"

                txt_text = f"[Файл name={filename} (содержимое: \"{txt_content}\")]"
                return txt_text + (f" {text}" if text else "")

    except Exception:
        pass

    return text


def _messages_to_dicts(response: object) -> list[dict[str, Any]]:
    """Extract and normalize the messages list from a LangGraph agent response."""
    if isinstance(response, Mapping):
        resp_map = cast("Mapping[str, Any]", response)
        raw_messages = resp_map.get("messages")
        if isinstance(raw_messages, list):
            return [_message_to_dict(m) for m in raw_messages]
    return []


def _message_to_dict(msg: object) -> dict[str, Any]:
    """Convert a LangChain BaseMessage (or dict) to a normalized dict."""
    if isinstance(msg, Mapping):
        return dict(cast("Mapping[str, Any]", msg))
    if hasattr(msg, "model_dump"):
        return msg.model_dump()  # type: ignore[no-any-return]
    # Fallback for plain objects with attributes
    result: dict[str, Any] = {}
    for attr in ("type", "role", "content", "tool_calls", "tool_call_id", "id", "name"):
        val = getattr(msg, attr, None)
        if val is not None:
            result[attr] = val
    return result


def _extract_response_text(response: object) -> str:
    if isinstance(response, str):
        return response
    if isinstance(response, Mapping):
        response_map = cast("Mapping[str, Any]", response)
        messages = response_map.get("messages")
        if isinstance(messages, list) and messages:
            # Find the last AIMessage (assistant response), not just the last message
            # (which could be a ToolMessage if tools were called).
            # BaseMessage objects use `type` ("ai", "human", "tool"), not `role`.
            for message in reversed(messages):
                if isinstance(message, Mapping):
                    msg_map = cast("Mapping[str, Any]", message)
                    role = msg_map.get("role")
                    if role == "assistant":
                        content = _get_content(message)
                        if content:
                            return content
                else:
                    # Check for BaseMessage objects (AIMessage has type="ai")
                    msg_type = getattr(message, "type", None)
                    role = getattr(message, "role", None)
                    if role == "assistant" or msg_type == "ai":
                        content = _get_content(message)
                        if content:
                            return content
        output = response_map.get("output") or response_map.get("content")
        if isinstance(output, str) and output:
            return output
    content = _get_content(response)
    if content:
        return content
    # Fallback: return empty string instead of str(response) to avoid sending garbage
    logger.warning("Could not extract response text from: %s", type(response))
    return ""


def _get_content(value: object) -> str | None:
    if isinstance(value, Mapping):
        value_map = cast("Mapping[str, Any]", value)
        content = value_map.get("content")
    else:
        content = getattr(value, "content", None)
    if isinstance(content, str) and content:
        return content
    if isinstance(content, list):
        parts: list[str] = []
        for item in content:
            if isinstance(item, str):
                parts.append(item)
            elif isinstance(item, Mapping):
                item_map = cast("Mapping[str, Any]", item)
                text = item_map.get("text")
                if isinstance(text, str):
                    parts.append(text)
        if parts:
            return "\n".join(parts)
    return None


def _interpret_agent_response(response: object) -> tuple[bool, str, int | None]:
    """Interpret a LangChain agent response into (send_any_message, text, reply_to).

    This function supports:
    - plain string responses
    - mapping-like responses produced by LangChain (dict-like), including parsed
      output from OutputParsers / PydanticOutputParser where fields like
      'send_any_message', 'text', and 'reply_to' may be present.

    Falls back to extracting the best available textual content.
    """
    # Default behaviour: send message, text from _extract_response_text, no reply
    default_text = _extract_response_text(response)

    if isinstance(response, Mapping):
        resp_map = cast("Mapping[str, Any]", response)
        # send_any_message may be absent; treat non-bool as True
        raw_send = resp_map.get("send_any_message")
        send_any = bool(raw_send) if isinstance(raw_send, bool) else True

        # prefer explicit 'text' field, then 'output'/'content' via extractor
        txt = resp_map.get("text")
        if isinstance(txt, str) and txt:
            text = txt
        else:
            text = default_text

        reply_val = resp_map.get("reply_to") or resp_map.get("reply_to_message_id")
        reply_to = None
        if isinstance(reply_val, int):
            reply_to = reply_val
        elif isinstance(reply_val, str) and reply_val.isdigit():
            try:
                reply_to = int(reply_val)
            except Exception:
                reply_to = None

        return send_any, text, reply_to

    # Non-mapping responses: send and use extracted text
    return True, default_text, None


def _extract_message_id(message: object) -> str | None:
    if isinstance(message, Mapping):
        message_map = cast("Mapping[str, Any]", message)
        value: Any = message_map.get("id")
    else:
        value = getattr(message, "id", None)
    if value is None:
        return None
    return str(value)


def _extract_incoming_text(event: object) -> str:
    text = getattr(event, "raw_text", None) or getattr(event, "text", None)
    if not isinstance(text, str):
        text = ""
        
    message = getattr(event, "message", None)
    if message and getattr(message, "media", None):
        try:
            from mimic42.integrations.telegram_tools import format_media_object
            media_id = format_media_object(message)
            if media_id:
                if media_id.startswith("photo:"):
                    text = f"[Фото id={media_id}]" + (f" {text}" if text else "")
                elif media_id.startswith("sticker:"):
                    parts = media_id.split(":")
                    emoji = parts[5] if len(parts) > 5 else ""
                    pack_name = parts[6] if len(parts) > 6 else ""
                    pack_str = f" пак={pack_name}" if pack_name else ""
                    text = (
                        f"[Стикер {emoji} id={media_id}{pack_str}]"
                        + (f" {text}" if text else "")
                    )
                elif media_id.startswith("voice:"):
                    text = (
                        f"[Голосовое сообщение id={media_id}]"
                        + (f" {text}" if text else "")
                    )
                elif media_id.startswith("round:"):
                    text = (
                        f"[Видеосообщение id={media_id}]"
                        + (f" {text}" if text else "")
                    )
                elif media_id.startswith("doc:"):
                    parts = media_id.split(":")
                    filename = parts[5] if len(parts) > 5 else "file"
                    text = (
                        f"[Файл name={filename} id={media_id}]"
                        + (f" {text}" if text else "")
                    )
        except Exception:
            pass
            
    return text


async def _extract_incoming_peer(event: object) -> str:
    chat_id = getattr(event, "chat_id", None)
    if chat_id is not None:
        return str(chat_id)
    get_chat = getattr(event, "get_chat", None)
    if callable(get_chat):
        chat = await get_chat()
        if chat:
            cid = getattr(chat, "id", None)
            if cid is not None:
                return str(cid)
    peer_id = getattr(event, "peer_id", None)
    if peer_id is not None:
        return str(peer_id)
    raise ValueError("Incoming Telegram event does not include a peer")


def _extract_incoming_message_id(event: object) -> int | None:
    value = getattr(event, "id", None)
    if isinstance(value, int):
        return value
    message = getattr(event, "message", None)
    message_id = getattr(message, "id", None)
    if isinstance(message_id, int):
        return message_id
    return None
